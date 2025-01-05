from docx import Document
from docx.shared import Pt
import os

# File Paths
TANAKH_DOCX_FOLDER = "tanakh_docs"
ENG_DOCX_FOLDER = "eng_docs"

def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def navigate_folders_from_cwd(start_path):
    """
    Dynamically navigate folders starting from the current working directory and return the final path.

    Returns:
    - str: The final folder path.
    """
    current_path = start_path

    while True:
        print(f"\nCurrent Path: {current_path}")
        subfolders = [f for f in os.listdir(current_path) if os.path.isdir(os.path.join(current_path, f))]

        if not subfolders:
            print("No subfolders available in the current directory.")
            break

        # Display subfolders with options
        print("Subfolders:")
        for i, folder in enumerate(subfolders, 1):
            print(f"{i}. {folder}")
        print("0. Select this folder")
        print("-1. Go up one level")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            # Select the current folder
            return current_path
        elif choice == -1:
            # Go up one level
            parent_path = os.path.dirname(current_path)
            if current_path == parent_path:  # At the root
                print("You are already at the root directory.")
            else:
                current_path = parent_path
        elif 1 <= choice <= len(subfolders):
            # Navigate into the selected subfolder
            current_path = os.path.join(current_path, subfolders[choice - 1])
        else:
            print("Invalid choice. Please try again.")

    return current_path

def pick_filename_from_folder(folder_path):
    """
    Dynamically pick a filename from the specified folder path.

    Parameters:
    - folder_path (str): The path to the folder.

    Returns:
    - str: The selected filename, or None if no valid file is selected.
    """
    if not os.path.isdir(folder_path):
        print("Invalid folder path. Please provide a valid directory.")
        return None

    while True:
        print(f"\nCurrent Folder: {folder_path}")
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

        if not files:
            print("No files available in this folder.")
            return None

        # Display files with options
        print("Files:")
        for i, file in enumerate(files, 1):
            print(f"{i}. {file}")
        print("0. Exit without selecting a file")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            print("No file selected.")
            return None
        elif 1 <= choice <= len(files):
            # Return the selected file
            selected_file = files[choice - 1]
            return os.path.join(folder_path, selected_file)
        else:
            print("Invalid choice. Please try again.")


def reformat_eng_docx(file_path):
    """
    Reformat the DOCX file so that each 'Verse' paragraph is on its own line,
    and other paragraphs under the same verse are merged into a single string (removing newlines).

    Args:
        file_path (str): Path to the DOCX file to reformat.
    """
    # Load the document
    doc = Document(file_path)

    # Create a list to hold the reformatted verses
    new_paragraphs = []
    current_verse = ""

    for para in doc.paragraphs:
        para_text = para.text.strip()

        # If the paragraph starts with "Verse", it should be the start of a new verse
        if para_text.startswith("Verse"):
            # If there's an accumulated verse, add it to the list
            if current_verse:
                new_paragraphs.append(current_verse)
            # Start a new verse and remove newlines from the text
            current_verse = para_text.replace("\n", " ")
        else:
            # Append non-"Verse" text to the current verse and remove any newlines
            if para_text:  # Avoid adding empty lines
                current_verse += " " + para_text.replace("\n", " ")

    # Add the last accumulated verse if it exists
    if current_verse:
        new_paragraphs.append(current_verse)

    # Create a new document and add the formatted paragraphs
    new_doc = Document()
    for para in new_paragraphs:
        new_doc.add_paragraph(para)

    # Save the new document with the same name, appending "_formatted"
    #formatted_file_path = file_path.replace(".docx", "_formatted.docx")
    
    # Remove the existing file if it already exists
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Existing file found and removed: {file_path}")
    
    new_doc.save(file_path)
    print(f"Formatted document saved as: {file_path}")


if __name__ == "__main__":
    # Example usage
    # Replace 'hebrew.docx', 'english.docx', and 'output.docx' with the actual file paths.
    eng_folder_path =load_tanakh_path(ENG_DOCX_FOLDER)
    english_file_path = "english_text.docx"

    final_eng_path = navigate_folders_from_cwd(eng_folder_path)
    print(f"\nEng - Final Path: {final_eng_path}")
    english_file = pick_filename_from_folder(final_eng_path)
    if english_file:
        print(f"\nSelected File: {english_file}")

    reformat_eng_docx(english_file)