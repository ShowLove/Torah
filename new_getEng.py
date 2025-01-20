# Import the homemade utils module
import utils

import os
import json
import time
import shutil
import subprocess
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def select_option(driver, dropdown_name, option_text):
    try:
        dropdown = Select(driver.find_element(By.NAME, dropdown_name))
        dropdown.select_by_visible_text(option_text)
        print(f"Option '{option_text}' selected from dropdown '{dropdown_name}'.")
    except Exception as e:
        print(f"An error occurred while selecting '{option_text}' from '{dropdown_name}': {e}")


def click_submit_button(driver):
    try:
        submit_button = driver.find_element(By.XPATH, "//input[@type='submit' and @value='GO']")
        submit_button.click()
        print("Submit button clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the submit button: {e}")


def get_current_url(driver):
    try:
        current_url = driver.current_url
        print(f"Current URL: {current_url}")
        return current_url
    except Exception as e:
        print(f"An error occurred while retrieving the current URL: {e}")
        return None


def grab_verses(driver):
    try:
        WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.TAG_NAME, "b")))
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, "html.parser")
        verses = soup.find_all('b')

        verse_texts = []

        for verse in verses:
            verse_number = verse.get_text(strip=True)
            verse_text = ""
            if verse.next_sibling:
                verse_text = verse.next_sibling.strip()
            else:
                next_p = verse.find_next('p')
                if next_p:
                    verse_text = next_p.get_text(strip=True)

            if verse_number and verse_text:
                verse_texts.append((verse_number, verse_text))

        return verse_texts

    except Exception as e:
        print(f"An error occurred while grabbing verses: {e}")
        return []


def save_to_word(verses, filename, book_name, chapter_number, file_path="."):
    os.makedirs(file_path, exist_ok=True)
    save_path = os.path.join(file_path, f"{filename}.docx")

    if os.path.exists(save_path):
        print(f"File exists, deleting: {save_path}")
        os.remove(save_path)

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = utils.MARGIN_SIZE
        section.right_margin = utils.MARGIN_SIZE
        section.top_margin = utils.MARGIN_SIZE
        section.bottom_margin = utils.MARGIN_SIZE

    doc.add_heading(f'{filename} - Chapter {chapter_number}', 0)

    for verse_number, verse_text in verses:
        doc.add_paragraph(f"{verse_number}: {verse_text}")

    doc.save(save_path)
    print(f"Saved Word document: {save_path}")


def get_Tanakh_and_verses(chapter_number, book_name):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://www.mnemotrix.com/texis/vtx/chumash")

    try:
        select_option(driver, "bookq", book_name)
        select_option(driver, "chapterq", f"Chapter {chapter_number}")
        time.sleep(3)
        click_submit_button(driver)
        time.sleep(2)

        partial_text = get_partial_text(book_name)
        links = driver.find_elements(By.PARTIAL_LINK_TEXT, partial_text)
        if links:
            links[0].click()
        else:
            print(f"----> No matching link found for {book_name} ch: {chapter_number}")
        time.sleep(2)

        final_url = get_current_url(driver)
        driver.get(final_url)
        verses = grab_verses(driver)

        filename = f"{book_name}_{chapter_number}.docx"
        folder_path = utils.load_tanakh_path(utils.ENG_DOCX_FOLDER)
        folder_path = os.path.join(folder_path, book_name)
        save_to_word(verses, filename, book_name, chapter_number, file_path=folder_path)

    finally:
        time.sleep(3)
        driver.quit()


def get_partial_text(book_name):
    book_mapping = {
        "Genesis": "Bereishis/Genesis, Chapter",
        "Exodus": "Shemos/Exodus, Chapter",
        "Leviticus": "Vayikro/Leviticus, Chapter",
        "Numbers": "Bamidbar/Numbers, Chapter",
        "Deuteronomy": "Devarim/Deuteronomy, Chapter",
    }
    return book_mapping.get(book_name, "Error: Invalid book input")


def process_specific_parasha(parasha_name, file_path="data/torah_parashot_eng.json"):
    with open(file_path, 'r') as file:
        parashot_data = json.load(file)

    for parasha in parashot_data["Parashot"]:
        if parasha["Name"] == parasha_name:
            parasha_book = parasha["Book"]
            start_chapter = int(parasha["Start"]["Chapter"])
            end_chapter = int(parasha["End"]["Chapter"])

            for current_chapter in range(start_chapter, end_chapter + 1):
                get_Tanakh_and_verses(str(current_chapter).zfill(2), parasha_book)
            return

    print(f"Parasha '{parasha_name}' not found in the file.")

def getChFromLink(parasha_link):

    # Get user inputs
    inputs = utils.get_tanakh_scraper_inputs()
    if not inputs:
        return
    tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice = inputs

    #Step 1 get driver
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get(parasha_link)  # Replace with your desired URL
    # Step 2 grab verses
    verses = grab_verses(driver)

    # Step 3: Save the verses to a Word document with the specified filename format
    filename = f"{book_name}_{chapter_choice}.docx"
    folder_path=load_tanakh_path(ENG_DOCX_FOLDER)
    #folder_path = os.path.join(folder_path, parasha_name)
    save_to_word(verses, filename, book_name, chapter_choice, file_path=folder_path)

def get_parasha_details(parasha_name, file_path="data/torah_parashot_eng.json"):
    # Load the JSON file
    with open(file_path, 'r') as file:
        parashot_data = json.load(file)
    
    # Search for the specific parasha
    for parasha in parashot_data["Parashot"]:
        if parasha["Name"] == parasha_name:
            # Return the parasha details
            return {
                "Name": parasha["Name"],
                "Book": parasha["Book"],
                "Start_Chapter": parasha["Start"]["Chapter"],
                "Start_Verse": parasha["Start"]["Verse"],
                "End_Chapter": parasha["End"]["Chapter"],
                "End_Verse": parasha["End"]["Verse"]
            }
    
    # If parasha not found
    print(f"Parasha '{parasha_name}' not found in the file.")
    return None

def main_open_website_with_chrome(website_url):
    try:
        # Path to the Google Chrome executable on macOS
        chrome_path = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        subprocess.Popen([chrome_path, website_url])  # Launch Chrome without tying it to the Python script
        print(f"Website {website_url} opened in Google Chrome successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
def prompt_user_choice():
    # Extract the first parasha's details into individual variables
    now_parasha_path = utils.load_data(utils.PARASHOT_NOW, return_path_only=True)
    details = utils.get_parasha_details(now_parasha_path)
    if details:
        first_parasha = details[0]
        Parasha = first_parasha["Parasha"]
        Book = first_parasha["Book"]
        Start = first_parasha["Start"]
        End = first_parasha["End"]
    else:
        print("No parasha details found.")
        return None

    # Ask the user to choose between the options
    print("Choose an option:")
    print("1. Open english Torah Site")
    print("2. Get parasha: " + Parasha)
    print("3. Get the chapter from a link")
    print("4. Get specific parasha details")

    choice = input("Please enter a number: 1 through 5.: ").strip()
    file_path = utils.load_data(utils.PARASHOT_LIST_ENG_FILE, return_path_only=True)

    if choice == "1":
        # Call the function to get all Genesis chapters from 1 to 50
        eng_website_url = "http://www.mnemotrix.com/texis/vtx/chumash"
        main_open_website_with_chrome(eng_website_url)
    elif choice == "2":
        process_specific_parasha(Parasha)
    elif choice == "3":
        link = "http://www.mnemotrix.com/texis/vtx/chumash/+3wwBme4J+he5z5xwwxFqhUw0GwqFqt0Ldm15mFqAgrwpBnGaX_nFqwtzmxwww/article.html"
        getChFromLink(link)
    if choice == "4":
        details = get_parasha_details("Vayechi")
        if details:
            print(details)
    else:
        print("Have a nice Day !")

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
if __name__ == "__main__":
    # Ask the user to choose between the options - retrieve data 
    prompt_user_choice()