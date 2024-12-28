import json                                                         # For encoding and decoding JSON data.                                                             
from selenium import webdriver                                      # For automating and controlling the web browser
from selenium.webdriver.common.by import By                         # For locating elements on the web page
from selenium.webdriver.chrome.service import Service               # For initializing and configuring the ChromeDriver
from webdriver_manager.chrome import ChromeDriverManager            # Manages downloading and setting up ChromeDriver
from selenium.webdriver.support.ui import Select                    # For interacting with drop-down menus (select elements)
import time                                                         # For pausing the execution of the program (e.g., sleep or wait)
import subprocess                                                   # For running system commands and interacting with the system shell
from bs4 import BeautifulSoup                                       # For parsing and navigating HTML or XML content
from selenium.webdriver.support.ui import WebDriverWait             # For waiting for elements to appear on the page
from selenium.webdriver.support import expected_conditions as EC    # For defining the expected conditions for elements
from docx import Document                                           # For creating and modifying Word documents
import os                                                           # For file and directory operations (e.g., working with paths, creating folders)
import shutil                                                       # For file operations (e.g., moving, copying, and deleting files)
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT                   #
from docx.oxml.ns import qn                                         #
from docx.oxml import OxmlElement                                   #
                                                                    ################################################################################################


# Global constants
DATA_FOLDER = "data"
TORAH_BOOKS = "Torah books"
PROPHETS_BOOKS = "Prophets books"
SCRIPTURES_BOOKS = "Scriptures books"
PENTATEUCH_FILE = "Pentateuch.json"
PROPHETS_FILE = "Prophets.json"
SCRIPTURES_FILE = "Scriptures.json"
TANAKH_OUTLINE_FILE = "tanakhOutlineHeb.json"
SCRAPER_URL = "https://www.chabad.org/library/bible_cdo/aid/63255/jewish/The-Bible-with-Rashi.htm"
TORAH_SECTION = "Torah (Pentateuch)"
PROPHETS_SECTION = "Nevi'im (Prophets)"
SCRIPTURES_SECTION = "Ketuvim (Scriptures)"

# Load data from the external JSON file
# Function to load JSON data from a file in the 'data' directory
def load_data(json_filename):
    file_path = os.path.join(DATA_FOLDER, json_filename)
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    else:
        print(f"Error: The file {json_filename} does not exist in the '{DATA_FOLDER}' folder.")
        return None

def prompt_user_for_book(data):
    print("Please choose a section:")
    for key, value in data['sections'].items():
        print(f"{key}. {value}")
    
    tanakh_divisions = input("Enter the number corresponding to your choice: ")

    if tanakh_divisions == "1":
        tanakh_division_name = "Torah books"
    elif tanakh_divisions == "2":
        tanakh_division_name = "Prophets books"
    elif tanakh_divisions == "3":
        tanakh_division_name = "Scriptures books"
    else:
        print("Invalid choice. Exiting...")
        return None, None, None
    
    print(f"\nYou selected: {data['sections'][tanakh_divisions]}")
    print("\nPlease choose a book:")

    books = data[tanakh_division_name]
    for key, value in books.items():
        print(f"{key}. {value}")

    book_choice = input("Enter the number corresponding to your choice: ")

    if book_choice in books:
        book_name = books[book_choice]
        print(f"\nYou selected: {book_name}")
        return tanakh_division_name, book_choice, book_name
    else:
        print("Invalid choice. Exiting...")
        return None, None, None

def is_valid_chapter(tanakh_division_name, book_choice, chapter_choice, verse_choice=None):
    if tanakh_division_name == TORAH_BOOKS:
        file_name = PENTATEUCH_FILE
    elif tanakh_division_name == PROPHETS_BOOKS:
        file_name = PROPHETS_FILE
    elif tanakh_division_name == SCRIPTURES_BOOKS:
        file_name = SCRIPTURES_FILE
    else:
        print("Invalid Tanakh division.")
        return False

    data = load_data(file_name)
    if data is None:
        return False

    if book_choice in data['books']:
        book_data = data['books'][book_choice]
    else:
        print("Invalid book choice.")
        return False

    if chapter_choice in book_data['chapters']:
        total_verses = book_data['chapters'][chapter_choice]
        
        if verse_choice is not None:
            if 1 <= verse_choice <= total_verses:
                return True
            else:
                print(f"Invalid verse choice. Chapter {chapter_choice} has {total_verses} verses.")
                return False
        else:
            return True
    else:
        print("Invalid chapter choice.")
        return False

def getTanakhBook():
    data = load_data("tanakhOutlineHeb.json")
    tanakh_division_name, book_choice_num, book_name = prompt_user_for_book(data)

    if not tanakh_division_name or not book_choice_num or not book_name:
        print("Exit: Invalid choice made")
        return None, None, None

    print(f"Tanakh Division: {tanakh_division_name}")
    print(f"Book Choice: {book_choice_num}")
    print(f"Book Name: {book_name}")

    return tanakh_division_name, book_choice_num, book_name

def get_chapter_and_verse_from_user(tanakh_division_name, book_name):
    # Prompt for chapter and verse input
    chapter_choice = input("Enter the chapter number: ")
    start_verse_choice = input("Enter the start verse number: ")
    end_verse_choice = input("Enter the end verse number: ")
    start_verse_choice = int(start_verse_choice) if start_verse_choice else None
    end_verse_choice = int(end_verse_choice) if end_verse_choice else None

    # Validate chapter and verse using the is_valid_chapter function
    is_valid = is_valid_chapter(tanakh_division_name, book_name, chapter_choice, end_verse_choice)

    if is_valid:
        print(f"Chapter {chapter_choice} and Verse {end_verse_choice if end_verse_choice else ''} are valid!")
        return chapter_choice, start_verse_choice, end_verse_choice  # Return the valid chapter and verse values
    else:
        print("Invalid chapter or verse choice.")
        return None, None, None  # Return None if invalid

def get_tanakh_scraper_inputs():
    """
    Handles user input for Tanakh scraping: book selection, chapter, and verse range.
    
    Returns:
        tuple: (tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice)
               or None if any input is invalid.
    """
    # Step 1: Choose the desired Book
    tanakh_division_name, book_choice_num, book_name = getTanakhBook()

    # Exit if invalid input
    if not tanakh_division_name or not book_choice_num or not book_name:
        print("Invalid book choice. Exiting...")
        return None

    # Step 2: Choose the chapter and verse range
    chapter_choice, start_verse_choice, end_verse_choice = get_chapter_and_verse_from_user(tanakh_division_name, book_name)
    if not chapter_choice or not start_verse_choice or not end_verse_choice:
        print("Invalid chapter or verse range. Exiting...")
        return None

    print(f"TANAKH: {tanakh_division_name}, {book_name}, {chapter_choice}:{start_verse_choice}-{end_verse_choice}")
    return tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice

def perform_tanakh_scraping(tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice):
    DEBUG = True  # Toggle for debug print statements
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get(SCRAPER_URL)

    try:
        if tanakh_division_name == TORAH_BOOKS:
            tanakh_division_name = TORAH_SECTION
        elif tanakh_division_name == PROPHETS_BOOKS:
            tanakh_division_name = PROPHETS_SECTION
        elif tanakh_division_name == SCRIPTURES_BOOKS:
            tanakh_division_name = SCRIPTURES_SECTION
        else:
            if DEBUG:
                print("Invalid choice. Exiting...")
            return

        select_option(driver, "Section", tanakh_division_name)
        select_option(driver, "Book", book_name)
        choose_chapter_with_driver(driver, chapter_choice)
        click_go_button(driver)
        click_close_button(driver)
        click_hebrew_toggle(driver)

        if DEBUG:
            print("Current website:", driver.current_url)

        verse_texts = get_verse_texts(driver, int(start_verse_choice), int(end_verse_choice))

        # Create a Word document with Hebrew-friendly formatting
        document = Document()
        document.add_heading(f"{book_name} - Chapter {chapter_choice} (Verses {start_verse_choice}-{end_verse_choice})", level=1)

        for verse_id, verse_text in verse_texts.items():
            paragraph = document.add_paragraph()
            paragraph.text = f"{verse_id}: {verse_text}"

            # Right-to-left alignment
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add Hebrew-specific styling
            run = paragraph.runs[0]
            run.font.name = "David"  # Use a Hebrew-friendly font

            # Ensure RTL is applied at the XML level
            paragraph._p.set(qn('w:bidi'), '1')

        # Define the file path
        os.makedirs("tanakh_docs", exist_ok=True)
        save_path = os.path.join(
            "tanakh_docs", 
            f"{book_name}_CH_{chapter_choice}_Verses_{start_verse_choice}_to_{end_verse_choice}.docx"
        )

        # Delete the file if it exists
        if os.path.exists(save_path):
            if DEBUG:
                print(f"File exists, deleting: {save_path}")
            os.remove(save_path)

        # Save the new document
        document.save(save_path)

        if DEBUG:
            print(f"Saved Hebrew-friendly Word document: {save_path}")

    except Exception as e:
        if DEBUG:
            print(f"An error occurred during scraping: {e}")
    finally:
        driver.quit()


##################################################################################
##################################################################################
# Web scraper functions start
##################################################################################
##################################################################################

##################################################################################
# Generic function to select an option from a dropdown
##################################################################################
def select_option(driver, dropdown_name, option_text):
    try:
        dropdown = Select(driver.find_element(By.NAME, dropdown_name))
        dropdown.select_by_visible_text(option_text)
        print(f"Option '{option_text}' selected from dropdown '{dropdown_name}'.")
    except Exception as e:
        print(f"Error selecting option '{option_text}' from dropdown '{dropdown_name}': {e}")

##################################################################################
# Function to click a specific link
##################################################################################
def click_link(driver, link_text):
    try:
        link = driver.find_element(By.LINK_TEXT, link_text)
        link.click()
        print(f"Link with text '{link_text}' clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the link '{link_text}': {e}")

##################################################################################
# Function to click a specific chapter
##################################################################################
def choose_chapter_with_driver(driver, chapter_choice):
    # Generate the chapter name dynamically
    chapter_name = "Chapter " + str(chapter_choice)
    print(f"Looking for: {chapter_name}")  # Debug
    
    # Wait for the dropdown to be available
    dropdown = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "autoListChapter"))
    )
    
    # Use Select class to interact with the dropdown
    select = Select(dropdown)
    
    # Iterate through options to find the desired chapter
    for option in select.options:
        print(f"Checking option: {repr(option.text.strip())}")  # Debug
        if chapter_name.lower() in option.text.strip().lower():
            select.select_by_visible_text(option.text)
            return None
    
    return None  # Return None if the chapter is not found

##################################################################################
# Click the go button
##################################################################################
def click_go_button(driver):
    try:
        # Wait for the "Go" button to be present
        go_button = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//span[text()='Go']"))
        )
        # Click the button
        go_button.click()
        print("Clicked the 'Go' button successfully.")
    except Exception as e:
        print(f"Failed to click the 'Go' button: {e}")

def click_close_button(driver):
    try:
        # Wait for the close button to be clickable
        close_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CLASS_NAME, "subscribe-popup-clmc__close-button"))
        )
        # Click the close button
        close_button.click()
        print("Clicked the close button successfully.")
    except Exception as e:
        print(f"Failed to click the close button: {e}")

def click_hebrew_toggle(driver):
    try:
        # Wait for the "Hebrew" input element to be present
        hebrew_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, "ToggleButtonsContainer_Hebrew"))
        )
        
        # Use JavaScript to click the element
        driver.execute_script("arguments[0].click();", hebrew_input)
        print("Successfully toggled the 'Hebrew' button.")
        return True
    except Exception as e:
        print(f"Failed to toggle the 'Hebrew' button: {e}")
        return False


def get_verse_texts(driver, N1, N):
    """
    Fetches the text of verses from N1 to N.

    Args:
        driver: Selenium WebDriver instance.
        N1 (int): Start verse number.
        N (int): End verse number.

    Returns:
        dict: A dictionary where keys are verse IDs (e.g., 'v1') and values are the verse texts.
    """
    verses = {}

    try:
        for i in range(N1, N + 1):
            verse_id = f"v{i}"
            # Locate the parent <td> with the class 'hebrew' and the specific ID
            verse_element = driver.find_element(By.CSS_SELECTOR, f"td.hebrew a[id='{verse_id}'] + span.co_VerseText")
            # Get the verse text
            verses[verse_id] = verse_element.text
    except Exception as e:
        print(f"Error occurred while fetching verses: {e}")

    return verses

##################################################################################
##################################################################################
# Web scraper functions end
##################################################################################
##################################################################################

def run_tanakh_scraper_main():

    # Get user inputs
    inputs = get_tanakh_scraper_inputs()
    if not inputs:
        return

    tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice = inputs

    # Perform the scraping process
    perform_tanakh_scraping(tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice)

# Example of how to call the function
if __name__ == "__main__":
    run_tanakh_scraper_main()

