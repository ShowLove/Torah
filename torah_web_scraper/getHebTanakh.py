import json                                                         # For encoding and decoding JSON data.                                                             
from selenium import webdriver                                      # For automating and controlling the web browser
from selenium.webdriver.common.by import By                         # For locating elements on the web page
from selenium.webdriver.chrome.service import Service               # For initializing and configuring the ChromeDriver
from webdriver_manager.chrome import ChromeDriverManager            # Manages downloading and setting up ChromeDriver
from selenium.webdriver.support.ui import Select                    # For interacting with drop-down menus (select elements)
import time                                                         # For pausing the execution of the program (e.g., sleep or wait)
import subprocess                                                   # For running system commands and interacting with the system shell
from bs4 import BeautifulSoup                                       # For parsing and navigating HTML or XML content
from selenium.webdriver.support.ui import WebDriverWait             # For waiting for elements to appear on the page
from selenium.webdriver.support import expected_conditions as EC    # For defining the expected conditions for elements
from docx import Document                                           # For creating and modifying Word documents
import os                                                           # For file and directory operations (e.g., working with paths, creating folders)
import shutil                                                       # For file operations (e.g., moving, copying, and deleting files)
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn 
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import inspect
                                                                    ################################################################################################


# Global constants
DATA_FOLDER = "data"
TANAKH_DOCX_FOLDER = "tanakh_docs"
HEB_DOCX_FOLDER = "hebrew_docs"
TORAH_BOOKS = "Torah books"
PROPHETS_BOOKS = "Prophets books"
SCRIPTURES_BOOKS = "Scriptures books"
PENTATEUCH_FILE = "Pentateuch.json"
PROPHETS_FILE = "Prophets.json"
SCRIPTURES_FILE = "Scriptures.json"
TANAKH_OUTLINE_FILE = "tanakhOutlineHeb.json"
PARASHOT_LIST_FILE = 'torah_parashot.json'
SCRAPER_URL = "https://www.chabad.org/library/bible_cdo/aid/63255/jewish/The-Bible-with-Rashi.htm"
TORAH_SECTION = "Torah (Pentateuch)"
PROPHETS_SECTION = "Nevi'im (Prophets)"
SCRIPTURES_SECTION = "Ketuvim (Scriptures)"
DOCX_HEBREW_FONT = "Frank Ruehl" # Use Frank Ruehl for Hebrew text on Word
FONT_SIZE = 18  # Font size in points
MARGIN_SIZE = Pt(18)  # Margin size in points
VERSE_ID_FONT_SIZE = 12  # Smaller font size for the verse ID

# Load data from the external JSON file
# Function to load JSON data from a file in the 'data' directory
def load_data(json_filename, return_path_only=False):
    file_path = os.path.join(DATA_FOLDER, json_filename)
    if return_path_only:
        return file_path
    
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    else:
        print(f"Error: The file {json_filename} does not exist in the '{DATA_FOLDER}' folder.")
        return None

# Load file path
def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def prompt_user_for_book(data):
    print("Please choose a section:")
    for key, value in data['sections'].items():
        print(f"{key}. {value}")
    
    tanakh_divisions = input("Enter the number corresponding to your choice: ")

    if tanakh_divisions == "1":
        tanakh_division_name = "Torah books"
    elif tanakh_divisions == "2":
        tanakh_division_name = "Prophets books"
    elif tanakh_divisions == "3":
        tanakh_division_name = "Scriptures books"
    else:
        print(f"Invalid choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None
    
    print(f"\nYou selected: {data['sections'][tanakh_divisions]}")
    print("\nPlease choose a book:")

    books = data[tanakh_division_name]
    for key, value in books.items():
        print(f"{key}. {value}")

    book_choice = input("Enter the number corresponding to your choice: ")

    if book_choice in books:
        book_name = books[book_choice]
        print(f"\nYou selected: {book_name}")
        return tanakh_division_name, book_choice, book_name
    else:
        print(f"Invalid choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None

def is_valid_chapter(tanakh_division_name, book_choice, chapter_choice, verse_choice=None):
    if tanakh_division_name == TORAH_BOOKS:
        file_name = PENTATEUCH_FILE
    elif tanakh_division_name == PROPHETS_BOOKS:
        file_name = PROPHETS_FILE
    elif tanakh_division_name == SCRIPTURES_BOOKS:
        file_name = SCRIPTURES_FILE
    else:
        print("Invalid Tanakh division.")
        return False

    data = load_data(file_name)
    if data is None:
        return False

    if book_choice in data['books']:
        book_data = data['books'][book_choice]
    else:
        print("Invalid book choice.")
        return False

    if chapter_choice in book_data['chapters']:
        total_verses = book_data['chapters'][chapter_choice]
        
        if verse_choice is not None:
            if 1 <= verse_choice <= total_verses:
                return True
            else:
                print(f"Invalid verse choice. Chapter {chapter_choice} has {total_verses} verses.")
                return False
        else:
            return True
    else:
        print("Invalid chapter choice.")
        return False

def getTanakhBook():
    data = load_data("tanakhOutlineHeb.json")
    tanakh_division_name, book_choice_num, book_name = prompt_user_for_book(data)

    if not tanakh_division_name or not book_choice_num or not book_name:
        print("Exit: Invalid choice made")
        return None, None, None

    print(f"Tanakh Division: {tanakh_division_name}")
    print(f"Book Choice: {book_choice_num}")
    print(f"Book Name: {book_name}")

    return tanakh_division_name, book_choice_num, book_name

def get_chapter_and_verse_from_user(tanakh_division_name, book_name):
    # Prompt for chapter and verse input
    chapter_choice = input("Enter the chapter number: ")
    start_verse_choice = input("Enter the start verse number: ")
    end_verse_choice = input("Enter the end verse number: ")
    start_verse_choice = int(start_verse_choice) if start_verse_choice else None
    end_verse_choice = int(end_verse_choice) if end_verse_choice else None

    # Validate chapter and verse using the is_valid_chapter function
    is_valid = is_valid_chapter(tanakh_division_name, book_name, chapter_choice, end_verse_choice)

    if is_valid:
        print(f"Chapter {chapter_choice} and Verse {end_verse_choice if end_verse_choice else ''} are valid!")
        return chapter_choice, start_verse_choice, end_verse_choice  # Return the valid chapter and verse values
    else:
        print("Invalid chapter or verse choice.")
        return None, None, None  # Return None if invalid

def get_tanakh_scraper_inputs(get_end_chapter=False):
    """
    Handles user input for Tanakh scraping: book selection, chapter, and verse range.
    
    Parameters:
        get_end_chapter (bool): Flag to enable input for end chapter choice.

    Returns:
        tuple: (tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, [end_chapter_choice if enabled])
               or None if any input is invalid.
    """
    # Step 1: Choose the desired Book
    tanakh_division_name, book_choice_num, book_name = getTanakhBook()

    # Exit if invalid input
    if not tanakh_division_name or not book_choice_num or not book_name:
        print("Invalid book choice. Exiting...")
        return None

    # Step 2: Choose the chapter and verse range
    chapter_choice, start_verse_choice, end_verse_choice = get_chapter_and_verse_from_user(tanakh_division_name, book_name)
    if not chapter_choice or not start_verse_choice or not end_verse_choice:
        print("Invalid chapter or verse range. Exiting...")
        return None

    # Step 3 (Optional): Get end chapter choice if flag is enabled
    end_chapter_choice = chapter_choice  # Default to the same chapter if not provided
    if get_end_chapter:
        try:
            end_chapter_choice = int(input("Enter the end chapter number: "))
            if int(end_chapter_choice) < int(chapter_choice):
                print("End chapter: {end_chapter_choice} cannot be less than the start chapter: {chapter_choice}. Exiting...")
                return None
        except ValueError:
            print("Invalid input for end chapter. Exiting...")
            return None

    # Print details for debugging
    print(f"TANAKH: {tanakh_division_name}, {book_name}, {chapter_choice}:{start_verse_choice}-{end_verse_choice}")
    if get_end_chapter:
        print(f"End Chapter: {end_chapter_choice}")

    # Return tuple including end_chapter_choice if enabled
    if get_end_chapter:
        return tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, end_chapter_choice
    return tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice

def number_to_hebrew(n):
    units = ["", "א", "ב", "ג", "ד", "ה", "ו", "ז", "ח", "ט"]
    tens = ["", "י", "כ", "ל", "מ", "נ", "ס", "ע", "פ", "צ"]
    hundreds = ["", "ק", "ר", "ש", "ת"]

    # Extract hundreds, tens, and units
    u = n % 10
    t = (n // 10) % 10
    h = (n // 100) % 10

    # Construct Hebrew number
    result = hundreds[h] + tens[t] + units[u]
    return result

def traverse_tanakh_scraper(tanakh_division_name, book_name=None, chapter_choice=None, end_chapter_choice=None, start_verse_choice=1, end_verse_choice=None, file_path=load_tanakh_path(HEB_DOCX_FOLDER)):
    """
    Traverses the Tanakh JSON structure and performs scraping for specified sections.
    """
    DEBUG = True  # Toggle for debug print statements

    try:
        # Load the JSON file
        with open(os.path.join("data", "Pentateuch.json"), 'r', encoding='utf-8') as file:
            tanakh_data = json.load(file)

        books = tanakh_data.get("books", {})
        
        # Traverse books in the division
        for current_book_name, book_data in books.items():
            if book_name and current_book_name != book_name:
                continue  # Skip if not the specified book
            
            chapters = book_data.get("chapters", {})
            
            for current_chapter, verse_count in chapters.items():
                current_chapter = int(current_chapter)  # Convert chapter to integer for comparison
                
                if chapter_choice and current_chapter < int(chapter_choice):
                    continue  # Skip chapters before the starting chapter
                
                if end_chapter_choice and current_chapter > int(end_chapter_choice):
                    break  # Stop processing beyond the ending chapter
                
                # Set the start and end verse ranges
                start_verse = start_verse_choice if current_chapter == chapter_choice else 1
                end_verse = end_verse_choice if current_chapter == end_chapter_choice else verse_count

                if DEBUG:
                    print(f"Processing {current_book_name}, Chapter {current_chapter}, Verses {start_verse}-{end_verse}")

                time.sleep(1) 
                # Perform scraping for the current range
                perform_tanakh_scraping(
                    tanakh_division_name=tanakh_division_name,
                    book_name=current_book_name,
                    chapter_choice=current_chapter,
                    start_verse_choice=start_verse,
                    end_verse_choice=end_verse,
                    file_path=file_path
                )
                
    except FileNotFoundError:
        print(f"Error: The file 'Pentateuch.json' was not found in the 'data' folder.")
    except json.JSONDecodeError:
        print("Error: Failed to decode JSON from 'Pentateuch.json'.")
    except Exception as e:
        if DEBUG:
            print(f"An unexpected error occurred: {e}")

##################################################################################
##################################################################################
# Web scraper functions start
##################################################################################
##################################################################################

def perform_tanakh_scraping(tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, file_path=load_tanakh_path(HEB_DOCX_FOLDER)):
    DEBUG = True  # Toggle for debug print statements
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get(SCRAPER_URL)

    try:
        # Map user-friendly names to scraper-specific names
        if tanakh_division_name == TORAH_BOOKS:
            tanakh_division_name = TORAH_SECTION
        elif tanakh_division_name == PROPHETS_BOOKS:
            tanakh_division_name = PROPHETS_SECTION
        elif tanakh_division_name == SCRIPTURES_BOOKS:
            tanakh_division_name = SCRIPTURES_SECTION
        else:
            if DEBUG:
                print(f"Invalid choice: {tanakh_division_name}, line: {inspect.currentframe().f_lineno}: Exiting...")
            return

        # Perform the scraping steps
        select_option(driver, "Section", tanakh_division_name)
        select_option(driver, "Book", book_name)
        choose_chapter_with_driver(driver, chapter_choice)
        click_go_button(driver)
        click_close_button(driver)
        click_hebrew_toggle(driver)

        if DEBUG:
            print("Current website:", driver.current_url)

        # Fetch verse texts
        verse_texts = get_verse_texts(driver, int(start_verse_choice), int(end_verse_choice))

        # Pass variables dynamically to create the Word document
        create_hebrew_word_document(book_name, chapter_choice, start_verse_choice, end_verse_choice, verse_texts, file_path=file_path)

    except Exception as e:
        if DEBUG:
            print(f"An error occurred during scraping: {e}")
    finally:
        driver.quit()


def create_hebrew_word_document(book_name, chapter_choice, start_verse_choice, end_verse_choice, verse_texts, file_path=load_tanakh_path(HEB_DOCX_FOLDER)):
    """
    Create a Word document with Hebrew-friendly formatting.
    """
    # Toggle to include or exclude verse IDs
    include_verse_id = True
    document = Document()

    # Set narrow margins
    sections = document.sections
    for section in sections:
        section.left_margin   = MARGIN_SIZE
        section.right_margin  = MARGIN_SIZE
        section.top_margin    = MARGIN_SIZE
        section.bottom_margin = MARGIN_SIZE

    # Add heading with dynamic values
    document.add_heading(
        f"{book_name} - Chapter {chapter_choice} (Verses {start_verse_choice}-{end_verse_choice})",
        level=1
    )

    for verse_id, verse_text in verse_texts.items():
        paragraph = document.add_paragraph()

        # Add the verse text with default styling
        run_text = paragraph.add_run(verse_text)
        run_text.font.name = DOCX_HEBREW_FONT
        run_text.font.size = Pt(FONT_SIZE)  # Standard font size for Hebrew text

        if include_verse_id:
            # Extract verse number by removing the "v" and convert to Hebrew
            verse_number = int(verse_id.lstrip('v'))  # Remove "v" and convert to integer
            verse_number_hebrew = number_to_hebrew(verse_number)  # Convert number to Hebrew
            
            # Add the Hebrew verse number at the beginning
            run_id = paragraph.add_run(f"  ({verse_number_hebrew})")
            run_id.font.name = DOCX_HEBREW_FONT
            run_id.font.size = Pt(FONT_SIZE)  # Smaller font size for the verse ID

        # Right-to-left alignment
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Ensure RTL is applied at the XML level
        paragraph._p.set(qn('w:bidi'), '1')

    # Define the file path dynamically
    os.makedirs(file_path, exist_ok=True)
    save_path = os.path.join(
        file_path,
        f"{book_name}_CH_{chapter_choice}_Verses_{start_verse_choice}_to_{end_verse_choice}.docx"
    )

    # Delete the file if it exists
    if os.path.exists(save_path):
        print(f"File exists, deleting: {save_path}")
        os.remove(save_path)

    # Save the new document
    document.save(save_path)
    print(f"Saved Hebrew-friendly Word document: {save_path}")

    # Do post processing
    docx_remove_colons(save_path, save_path)
    docx_add_colons(save_path, save_path)

def docx_remove_colons(input_path, output_path):

    # Removes all colons (:) from a Word document while preserving the formatting.

    # Load the document
    doc = Document(input_path)
    
    # Remove all the ":" from the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(":", "")  # Replace ":" in each run's text
    
    # Remove all the ":" from the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(":", "")  # Replace ":" in table cells
    
    # Save the modified document
    doc.save(output_path)

def docx_add_colons(input_path, output_path):
    
    # Adds a colon at the end of each Hebrew sentence in a Word document, ensuring proper placement for right-to-left text.

    # Load the document
    doc = Document(input_path)

    # Process paragraphs
    for paragraph in doc.paragraphs:
        # Set paragraph alignment to Right-to-Left
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        for run in paragraph.runs:
            # Strip trailing spaces
            text = run.text.strip()
            
            # Add a colon if the text doesn't already end with one
            if text and not text.endswith(":"):
                text += "‪:"  # Add a colon with an RTL marker
            
            # Apply Hebrew font and formatting
            run.text = text
            run.font.name = "David"  # Example Hebrew font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "David")
    
    # Save the modified document
    doc.save(output_path)

##################################################################################
# Generic function to select an option from a dropdown
##################################################################################
def select_option(driver, dropdown_name, option_text):
    try:
        dropdown = Select(driver.find_element(By.NAME, dropdown_name))
        dropdown.select_by_visible_text(option_text)
        print(f"Option '{option_text}' selected from dropdown '{dropdown_name}'.")
    except Exception as e:
        print(f"Error selecting option '{option_text}' from dropdown '{dropdown_name}': {e}")

##################################################################################
# Function to click a specific link
##################################################################################
def click_link(driver, link_text):
    try:
        link = driver.find_element(By.LINK_TEXT, link_text)
        link.click()
        print(f"Link with text '{link_text}' clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the link '{link_text}': {e}")

##################################################################################
# Function to click a specific chapter
##################################################################################
def choose_chapter_with_driver(driver, chapter_choice):
    # Generate the chapter name dynamically
    chapter_name = "Chapter " + str(chapter_choice)
    print(f"Looking for: {chapter_name}")  # Debug
    
    # Wait for the dropdown to be available
    dropdown = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located((By.ID, "autoListChapter"))
    )
    
    # Use Select class to interact with the dropdown
    select = Select(dropdown)
    
    # Iterate through options to find the desired chapter
    for option in select.options:
        print(f"Checking option: {repr(option.text.strip())}")  # Debug
        if chapter_name.lower() in option.text.strip().lower():
            select.select_by_visible_text(option.text)
            return None
    
    return None  # Return None if the chapter is not found

##################################################################################
# Click the go button
##################################################################################
def click_go_button(driver):
    try:
        # Wait for the "Go" button to be present
        go_button = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//span[text()='Go']"))
        )
        # Click the button
        go_button.click()
        print("Clicked the 'Go' button successfully.")
    except Exception as e:
        print(f"Failed to click the 'Go' button: {e}")

def click_close_button(driver):
    try:
        # Wait for the close button to be clickable
        close_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CLASS_NAME, "subscribe-popup-clmc__close-button"))
        )
        # Click the close button
        close_button.click()
        print("Clicked the close button successfully.")
    except Exception as e:
        print(f"Failed to click the close button: {e}")

def click_hebrew_toggle(driver):
    try:
        # Wait for the "Hebrew" input element to be present
        hebrew_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, "ToggleButtonsContainer_Hebrew"))
        )
        
        # Use JavaScript to click the element
        driver.execute_script("arguments[0].click();", hebrew_input)
        print("Successfully toggled the 'Hebrew' button.")
        return True
    except Exception as e:
        print(f"Failed to toggle the 'Hebrew' button: {e}")
        return False


def get_verse_texts(driver, N1, N):
    """
    Fetches the text of verses from N1 to N.

    Args:
        driver: Selenium WebDriver instance.
        N1 (int): Start verse number.
        N (int): End verse number.

    Returns:
        dict: A dictionary where keys are verse IDs (e.g., 'v1') and values are the verse texts.
    """
    verses = {}

    try:
        for i in range(N1, N + 1):
            verse_id = f"v{i}"
            # Locate the parent <td> with the class 'hebrew' and the specific ID
            verse_element = driver.find_element(By.CSS_SELECTOR, f"td.hebrew a[id='{verse_id}'] + span.co_VerseText")
            # Get the verse text
            verses[verse_id] = verse_element.text
    except Exception as e:
        print(f"Error occurred while fetching verses: {e}")

    return verses

##################################################################################
##################################################################################
# Web scraper functions end
##################################################################################
##################################################################################

def run_tanakh_scraper_main():

    # Get user inputs
    inputs = get_tanakh_scraper_inputs()
    if not inputs:
        return

    tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice = inputs

    # Perform the scraping process
    perform_tanakh_scraping(tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice)

def print_parashah_info_main(file_name):
    """
    This function loads the JSON file containing Parashot information and prints the details
    (Name, Book, Tanakh Section, Start, and End) for each Parashah.

    :param file_name: Path to the JSON file containing Torah Parashot data.
    """
    try:
        # Load the JSON data from the file
        with open(file_name, 'r') as file:
            data = json.load(file)

        # Traverse through the Parashot list and print the details
        for parashah in data['Parashot']:
            name = parashah['Name']
            book = parashah['Book']
            tanakh_section = parashah['Tanakh Section']
            start_chapter = parashah['Start']['Chapter']
            start_verse = parashah['Start']['Verse']
            end_chapter = parashah['End']['Chapter']
            end_verse = parashah['End']['Verse']
            
            # Print the information in a readable format
            print(f"Parashah: {name}")
            print(f"Book: {book}")
            print(f"Tanakh Section: {tanakh_section}")
            print(f"Start: Chapter {start_chapter}, Verse {start_verse}")
            print(f"End: Chapter {end_chapter}, Verse {end_verse}")
            print("-" * 40)  # Separator for clarity

    except FileNotFoundError:
        print(f"Error: The file '{file_name}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: There was an issue decoding the JSON data in the file '{file_name}'.")
    except KeyError as e:
        print(f"Error: Missing expected key {e} in the JSON data.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def get_tanakh_range_from_input_main():

    inputs = get_tanakh_scraper_inputs(get_end_chapter=True)
    if not inputs:
        print("Error: no inputs in get_tanakh_scraper_inputs")

    tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, end_chapter_choice = inputs

    # Traverse and scrape Genesis
    traverse_tanakh_scraper(
        tanakh_division_name=tanakh_division_name,
        book_name=book_name,
        chapter_choice=chapter_choice,
        end_chapter_choice=end_chapter_choice,
        start_verse_choice=start_verse_choice,
        end_verse_choice=end_verse_choice
    )

def get_tanakh_range_from_json_main(parasha_name, file_path="data/torah_parashot.json"):
    """
    Gets the Tanakh range based on the specified parasha name from torah_parashot.json.

    :param parasha_name: Name of the Torah portion (e.g., "Bereshit", "Noach").
    :param file_path: Path to the torah_parashot.json file.
    """
    try:
        # Load the parasha data from the JSON file
        with open(file_path, 'r', encoding='utf-8') as file:
            parasha_data = json.load(file)
        
        # Find the specified parasha
        parasha = next((p for p in parasha_data["Parashot"] if p["Name"] == parasha_name), None)
        
        if not parasha:
            print(f"Error: Parasha '{parasha_name}' not found in the data.")
            return
        
        # Extract range information
        tanakh_division_name = parasha["Tanakh Section"]
        book_name = parasha["Book"]
        start_chapter = parasha["Start"]["Chapter"]
        start_verse = parasha["Start"]["Verse"]
        end_chapter = parasha["End"]["Chapter"]
        end_verse = parasha["End"]["Verse"]

        if tanakh_division_name == "Torah (Pentateuch)":
            tanakh_division_name = "Torah books"
        elif tanakh_divisions == "Nevi'im (Prophets)":
            tanakh_division_name = "Prophets books"
        elif tanakh_divisions == "Ketuvim (Scriptures)":
            tanakh_division_name = "Scriptures books"
        else:
            print(f"Invalid choice. line: {inspect.currentframe().f_lineno}: Exiting...")
            return None, None, None
        
        # Traverse and scrape using the extracted range
        folder_path=load_tanakh_path(HEB_DOCX_FOLDER)
        folder_path = os.path.join(folder_path, parasha_name)
        traverse_tanakh_scraper(
            tanakh_division_name=tanakh_division_name,
            book_name=book_name,
            chapter_choice=start_chapter,
            end_chapter_choice=end_chapter,
            start_verse_choice=start_verse,
            end_verse_choice=end_verse,
            file_path=folder_path
        )
    
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: Failed to decode JSON from '{file_path}'.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def process_all_parashot_main(file_path="data/torah_parashot.json"):
    """
    Processes all the parashot in the torah_parashot.json file by passing their names to
    the get_tanakh_range_from_json_main function.

    :param file_path: Path to the torah_parashot.json file.
    """
    try:
        # Load the parasha data from the JSON file
        with open(file_path, 'r', encoding='utf-8') as file:
            parasha_data = json.load(file)
        
        # Loop through all the "Name" fields and process each parasha
        for parasha in parasha_data.get("Parashot", []):
            parasha_name = parasha.get("Name")
            if parasha_name:
                print(f"Processing parasha: {parasha_name}")
                time.sleep(20)
                get_tanakh_range_from_json_main(parasha_name, file_path)
            else:
                print("Skipping a parasha with missing 'Name' field.")
    
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: Failed to decode JSON from '{file_path}'.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def rename_folders_by_timestamp_main(directory_path):
    try:
        # Get all folders in the directory
        folders = [f for f in os.listdir(directory_path) if os.path.isdir(os.path.join(directory_path, f))]
        
        # Get folders with their creation timestamps
        folders_with_timestamps = [(folder, os.path.getctime(os.path.join(directory_path, folder))) for folder in folders]
        
        # Sort folders by timestamp (oldest first)
        folders_with_timestamps.sort(key=lambda x: x[1])
        
        # Rename folders with prefix numbers
        for i, (folder, _) in enumerate(folders_with_timestamps, start=1):
            old_path = os.path.join(directory_path, folder)
            new_folder_name = f"{i:02d}_{folder}"  # Add zero-padded number as prefix
            new_path = os.path.join(directory_path, new_folder_name)
            os.rename(old_path, new_path)
            print(f"Renamed '{folder}' to '{new_folder_name}'")
    
    except Exception as e:
        print(f"An error occurred: {e}")

# Example of how to call the function
if __name__ == "__main__":
    # Prompt the user to choose the function to run
    print("Choose an option:")
    print("1. Run Tanakh Scraper")
    print("2. Print Parashah Info")
    print("3. Traverse Tanakh Scraper from a start to an end point")
    print("4. Get all the hebrew text from a single parasha portion")
    print("5. Get all the Parashot")
    print("6. Define folder names in a directory by time-stamp")

    # Get user input
    choice = input("Enter 1 through 6: ")
    file_name = load_data(PARASHOT_LIST_FILE, return_path_only=True)

    if choice == '1':
        # Call the Tanakh scraper function
        run_tanakh_scraper_main()
    elif choice == '2':
        # Call the Parashah info printing function
        print_parashah_info_main(file_name)
    elif choice == '3':
        # Scrape range of tanakh from user input
        get_tanakh_range_from_input_main()
    elif choice == '4':
        # Scrape the docx for a single parasha
        print_parashah_info_main(file_name)
        parasha_choice = input("Enter a parasha choice: ")
        get_tanakh_range_from_json_main(parasha_choice)
    elif choice == '5':
        process_all_parashot_main("data/torah_parashot.json")
    elif choice == '6':
        rename_folders_by_timestamp_main("tanakh_docs/hebrew_docs")
    else:
        print("Invalid choice. Please enter 1 through 5.")