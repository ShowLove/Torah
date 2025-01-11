from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import re
import os
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# File Paths
TANAKH_DOCX_FOLDER = "tanakh_docs"
ENG_DOCX_FOLDER = "eng_docs"
HEB_DOCX_FOLDER = "hebrew_docs"
OUTPUT_DOCX_FOLDER = "output_docs"

# Configurable constants for formatting
DOCX_HEBREW_FONT = "Frank Ruehl"  # Use Frank Ruehl for Hebrew text on Word
DOCX_ENGLISH_FONT = "Times New Roman"  # Use Times New Roman for English text
FONT_SIZE_HEB = 16  # Font size in points
FONT_SIZE_ENG = 12  # Font size in points
MARGIN_SIZE = Pt(12)  # Margin size in points

def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def pick_filename_from_folder(folder_path):
    """
    Dynamically pick a filename from the specified folder path.

    Parameters:
    - folder_path (str): The path to the folder.

    Returns:
    - str: The selected filename, or None if no valid file is selected.
    """
    if not os.path.isdir(folder_path):
        print("Invalid folder path. Please provide a valid directory.")
        return None

    while True:
        print(f"\nCurrent Folder: {folder_path}")
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

        if not files:
            print("No files available in this folder.")
            return None

        # Display files with options
        print("Files:")
        for i, file in enumerate(files, 1):
            print(f"{i}. {file}")
        print("0. Exit without selecting a file")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            print("No file selected.")
            return None
        elif 1 <= choice <= len(files):
            # Return the selected file
            selected_file = files[choice - 1]
            return os.path.join(folder_path, selected_file)
        else:
            print("Invalid choice. Please try again.")

def remove_second_colon_eng(para_text):
    """
    Replace all occurrences of "::" with ":" in the given text.

    Parameters:
    - para_text (str): The text of the paragraph.

    Returns:
    - str: The updated text with "::" replaced by ":".
    """
    # Replace all occurrences of "::" with ":"
    return re.sub(r"::", r":", para_text)

def process_paragraph(paragraph):
    """
    Process the paragraph to replace "::" with ":" while preserving formatting.

    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph object to process.
    """
    # Combine all runs' text into a single string
    para_text = "".join(run.text for run in paragraph.runs)

    # Replace "::" with ":" in the combined text
    updated_text = remove_second_colon_eng(para_text)

    # If changes are needed, clear existing runs and re-add the updated text
    if para_text != updated_text:
        for run in paragraph.runs:
            run.text = ""  # Clear existing text
        paragraph.add_run(updated_text)  # Add updated text back

def update_second_line(document):
    """
    Update the second line of the document by removing everything before 'Chapter'.

    Parameters:
    - document (docx.Document): The document object to process.
    """
    # Ensure the document has at least two paragraphs
    if len(document.paragraphs) >= 2:
        second_paragraph = document.paragraphs[1]
        if "Chapter" in second_paragraph.text:
            # Find the position of the word "Chapter" and update the text
            chapter_index = second_paragraph.text.find("Chapter")
            updated_text = second_paragraph.text[chapter_index:]
            second_paragraph.text = updated_text

def format_hebrew_paragraph(document):
    """
    Format Hebrew text in a Word document:
    - Align the text to the right (RTL).
    - Move the verse number to the beginning of the paragraph.
    - Apply specific fonts and font sizes.

    Parameters:
    - document (docx.Document): The document to be formatted.
    """
    
    for paragraph in document.paragraphs:
        # Step 1: Check if the paragraph contains any Hebrew characters
        if contains_hebrew(paragraph.text):
            # Step 2: Apply right-to-left (RTL) alignment
            apply_rtl_alignment(paragraph)
            
            # Step 3: Move the verse number (if present) to the beginning of the paragraph
            paragraph.text = move_verse_number_to_start(paragraph.text)
            
            # Step 4: Apply the Hebrew font and size to the paragraph
            apply_hebrew_font(paragraph)

def contains_hebrew(text):
    """
    Check if the provided text contains Hebrew characters.

    Parameters:
    - text (str): The text to check.

    Returns:
    - bool: True if the text contains Hebrew characters, False otherwise.
    """
    # Hebrew characters range from Unicode 0x0590 to 0x05FF
    return any("\u0590" <= char <= "\u05FF" for char in text)

def apply_rtl_alignment(paragraph):
    """
    Apply right-to-left alignment to the paragraph and ensure it is set at the XML level.
    
    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph to format.
    """
    # Align paragraph text to the right
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Ensure the XML-level setting for RTL is applied
    paragraph._p.set(qn('w:bidi'), '1')

def move_verse_number_to_start(text):
    """
    Move the verse number (if present) to the start of the paragraph.
    
    Parameters:
    - text (str): The paragraph text.
    
    Returns:
    - str: The modified text with the verse number moved to the beginning.
    """
    print(f"Original text: {text}")  # Debug: Show the original text
    
    # Regular expression to match verse numbers inside parentheses, including special characters like ‪
    match = re.search(r"‪\s?\(([^)]+)\)\s?‪", text)  # Match "(כח)", "(כט)", "(לא)", etc. allowing for spaces or special characters
    
    if match:
        # Extract the verse number (e.g., "כח", "כט")
        verse_number = match.group(1)
        print(f"Verse number found: {verse_number}")  # Debug: Show the found verse number
        
        # Remove the verse number from its original position
        modified_text = re.sub(r"‪\s?\(.*?\)\s?‪", "", text).strip()  # Strip any unwanted spaces
        
        # Remove any trailing colon (:) from the text
        if modified_text.endswith(":"):
            modified_text = modified_text[:-1]
        
        # Build the new text with the verse number at the start
        modified_text = modified_text + "\u200F"
        verse_number = verse_number + "\u200F"
        new_text =  f" {verse_number} " + modified_text + " " + ":\u200F"
        print(f"New text after modification: {new_text}")  # Debug: Show the modified text
        
        return new_text
    else:
        print("No verse number found.")  # Debug: Inform when no verse number is found
        # If no verse number is found, return the original text
        return text

def apply_hebrew_font(paragraph):
    """
    Apply the specific Hebrew font and size to the paragraph text.
    
    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph to format.
    """
    for run in paragraph.runs:
        # Apply Hebrew font and size
        run.font.name = DOCX_HEBREW_FONT
        run.font.size = Pt(FONT_SIZE_HEB)

if __name__ == "__main__":
    eng_folder_path = load_tanakh_path(ENG_DOCX_FOLDER)
    heb_folder_path = load_tanakh_path(HEB_DOCX_FOLDER)
    output_folder_path = load_tanakh_path(OUTPUT_DOCX_FOLDER)

    # Step 3: Do post-processing
    final_output_file = pick_filename_from_folder(output_folder_path)
    if final_output_file:
        print(f"\nSelected File: {final_output_file}")

    doc = Document(final_output_file)

    # Update the header
    update_second_line(doc)

    # Format the Hebrew text
    format_hebrew_paragraph(doc)

    # Iterate over paragraphs and process each one
    for para in doc.paragraphs:
        process_paragraph(para)

    # Save the updated document
    doc.save(final_output_file)
    print("File has been updated and saved.")