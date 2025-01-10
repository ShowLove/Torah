from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import re
import os
from docx.shared import RGBColor

# File Paths
TANAKH_DOCX_FOLDER = "tanakh_docs"
ENG_DOCX_FOLDER = "eng_docs"
HEB_DOCX_FOLDER = "hebrew_docs"
OUTPUT_DOCX_FOLDER = "output_docs"

# Configurable constants for formatting
DOCX_HEBREW_FONT = "Frank Ruehl"  # Use Frank Ruehl for Hebrew text on Word
DOCX_ENGLISH_FONT = "Times New Roman"  # Use Times New Roman for English text
FONT_SIZE_HEB = 16  # Font size in points
FONT_SIZE_ENG = 12  # Font size in points
MARGIN_SIZE = Pt(12)  # Margin size in points

def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def pick_filename_from_folder(folder_path):
    """
    Dynamically pick a filename from the specified folder path.

    Parameters:
    - folder_path (str): The path to the folder.

    Returns:
    - str: The selected filename, or None if no valid file is selected.
    """
    if not os.path.isdir(folder_path):
        print("Invalid folder path. Please provide a valid directory.")
        return None

    while True:
        print(f"\nCurrent Folder: {folder_path}")
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

        if not files:
            print("No files available in this folder.")
            return None

        # Display files with options
        print("Files:")
        for i, file in enumerate(files, 1):
            print(f"{i}. {file}")
        print("0. Exit without selecting a file")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            print("No file selected.")
            return None
        elif 1 <= choice <= len(files):
            # Return the selected file
            selected_file = files[choice - 1]
            return os.path.join(folder_path, selected_file)
        else:
            print("Invalid choice. Please try again.")

def remove_second_colon_eng(para_text):
    """
    Replace all occurrences of "::" with ":" in the given text.

    Parameters:
    - para_text (str): The text of the paragraph.

    Returns:
    - str: The updated text with "::" replaced by ":".
    """
    # Replace all occurrences of "::" with ":"
    return re.sub(r"::", r":", para_text)

def process_paragraph(paragraph):
    """
    Process the paragraph to replace "::" with ":" while preserving formatting.

    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph object to process.
    """
    # Combine all runs' text into a single string
    para_text = "".join(run.text for run in paragraph.runs)

    # Replace "::" with ":" in the combined text
    updated_text = remove_second_colon_eng(para_text)

    # If changes are needed, clear existing runs and re-add the updated text
    if para_text != updated_text:
        for run in paragraph.runs:
            run.text = ""  # Clear existing text
        paragraph.add_run(updated_text)  # Add updated text back

if __name__ == "__main__":
    eng_folder_path = load_tanakh_path(ENG_DOCX_FOLDER)
    heb_folder_path = load_tanakh_path(HEB_DOCX_FOLDER)
    output_folder_path = load_tanakh_path(OUTPUT_DOCX_FOLDER)

    # Step 3 do post processing
    final_output_file = pick_filename_from_folder(output_folder_path)
    if final_output_file:
        print(f"\nSelected File: {final_output_file}")

    #add_notes_to_verses(final_output_file)
    doc = Document(final_output_file)

    # Iterate over paragraphs and process each one
    for para in doc.paragraphs:
        process_paragraph(para)

    # Save the updated document
    doc.save(final_output_file)
    print("File has been updated and saved.")