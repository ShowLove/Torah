from selenium import webdriver                                      # For automating and controlling the web browser
from selenium.webdriver.common.by import By                         # For locating elements on the web page
from selenium.webdriver.chrome.service import Service               # For initializing and configuring the ChromeDriver
from webdriver_manager.chrome import ChromeDriverManager            # Manages downloading and setting up ChromeDriver
from selenium.webdriver.support.ui import Select                    # For interacting with drop-down menus (select elements)
import time                                                         # For pausing the execution of the program (e.g., sleep or wait)
import subprocess                                                   # For running system commands and interacting with the system shell
from bs4 import BeautifulSoup                                       # For parsing and navigating HTML or XML content
from selenium.webdriver.support.ui import WebDriverWait             # For waiting for elements to appear on the page
from selenium.webdriver.support import expected_conditions as EC    # For defining the expected conditions for elements
from docx import Document                                           # For creating and modifying Word documents
import os                                                           # For file and directory operations (e.g., working with paths, creating folders)
import shutil                                                       # For file operations (e.g., moving, copying, and deleting files)
import json
                                                                    ################################################################################################

TANAKH_OUTLINE_ENG = "tanakhOutlineEng.json"
DATA_FOLDER = "data"
PENTATEUCH_FILE_ENG = "Pentateuch_eng.json"
PROPHETS_FILE_ENG = "Prophets_eng.json"
SCRIPTURES_FILE_ENG = "Scriptures_eng.json"
TORAH_BOOKS = "Torah books"
PROPHETS_BOOKS = "Prophets books"
SCRIPTURES_BOOKS = "Scriptures books"
TANAKH_DOCX_FOLDER = "tanakh_docs"
ENG_DOCX_FOLDER = "eng_docs"
PARASHOT_LIST_ENG_FILE = 'torah_parashot_eng.json'

# Load data from the external JSON file
# Function to load JSON data from a file in the 'data' directory
def load_data(json_filename, return_path_only=False):
    file_path = os.path.join(DATA_FOLDER, json_filename)
    if return_path_only:
        return file_path
    
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    else:
        print(f"Error: The file {json_filename} does not exist in the '{DATA_FOLDER}' folder.")
        return None

def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def print_parashah_info_main(file_name):
    """
    This function loads the JSON file containing Parashot information and prints the details
    (Name, Book, Tanakh Section, Start, and End) for each Parashah.

    :param file_name: Path to the JSON file containing Torah Parashot data.
    """
    try:
        # Load the JSON data from the file
        with open(file_name, 'r') as file:
            data = json.load(file)

        # Traverse through the Parashot list and print the details
        for parashah in data['Parashot']:
            name = parashah['Name']
            book = parashah['Book']
            tanakh_section = parashah['Tanakh Section']
            start_chapter = parashah['Start']['Chapter']
            start_verse = parashah['Start']['Verse']
            end_chapter = parashah['End']['Chapter']
            end_verse = parashah['End']['Verse']
            
            # Print the information in a readable format
            print(f"Parashah: {name}")
            print(f"Book: {book}")
            print(f"Tanakh Section: {tanakh_section}")
            print(f"Start: Chapter {start_chapter}, Verse {start_verse}")
            print(f"End: Chapter {end_chapter}, Verse {end_verse}")
            print("-" * 40)  # Separator for clarity

    except FileNotFoundError:
        print(f"Error: The file '{file_name}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: There was an issue decoding the JSON data in the file '{file_name}'.")
    except KeyError as e:
        print(f"Error: Missing expected key {e} in the JSON data.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def get_tanakh_scraper_inputs(get_end_chapter=False):
    """
    Handles user input for Tanakh scraping: book selection, chapter, and verse range.
    
    Parameters:
        get_end_chapter (bool): Flag to enable input for end chapter choice.

    Returns:
        tuple: (tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, [end_chapter_choice if enabled])
               or None if any input is invalid.
    """
    DEBUG = True  # Toggle for debug print statements

    # Step 1: Choose the desired Book
    tanakh_division_name, book_choice_num, book_name = getTanakhBook()

    # Exit if invalid input
    if not tanakh_division_name or not book_choice_num or not book_name:
        print(f"Invalid book choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None

    # Step 2: Choose the chapter and verse range
    chapter_choice, start_verse_choice, end_verse_choice = get_chapter_and_verse_from_user(tanakh_division_name, book_name)
    if not chapter_choice or not start_verse_choice or not end_verse_choice:
        print(f"Invalid chapter or verse range. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None

    # Step 3 (Optional): Get end chapter choice if flag is enabled
    end_chapter_choice = chapter_choice  # Default to the same chapter if not provided
    if get_end_chapter:
        try:
            end_chapter_choice = int(input("Enter the end chapter number: "))
            if int(end_chapter_choice) < int(chapter_choice):
                print("End chapter: {end_chapter_choice} cannot be less than the start chapter: {chapter_choice}. Exiting...")
                return None
        except ValueError:
            print(f"Invalid input for end chapter. line: {inspect.currentframe().f_lineno}: Exiting...")
            return None

    # Print details for debugging
    if DEBUG:
        print(f"TANAKH: {tanakh_division_name}, {book_name}, {chapter_choice}:{start_verse_choice}-{end_verse_choice}")
        if get_end_chapter:
            print(f"End Chapter: {end_chapter_choice}")

    # Return tuple including end_chapter_choice if enabled
    if get_end_chapter:
        return tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice, end_chapter_choice
    return tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice


def getTanakhBook():

    DEBUG = True  # Toggle for debug print statements

    data = load_data(TANAKH_OUTLINE_ENG)
    tanakh_division_name, book_choice_num, book_name = prompt_user_for_book(data)

    if not tanakh_division_name or not book_choice_num or not book_name:
        print(f"Exit: Invalid choice made. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None

    if DEBUG:
        print(f"Tanakh Division: {tanakh_division_name}")
        print(f"Book Choice: {book_choice_num}")
        print(f"Book Name: {book_name}")

    return tanakh_division_name, book_choice_num, book_name

def get_chapter_and_verse_from_user(tanakh_division_name, book_name):
    # Prompt for chapter and verse input
    chapter_choice = input("Enter the chapter number: ")
    start_verse_choice = input("Enter the start verse number: ")
    end_verse_choice = input("Enter the end verse number: ")
    start_verse_choice = int(start_verse_choice) if start_verse_choice else None
    end_verse_choice = int(end_verse_choice) if end_verse_choice else None

    # Validate chapter and verse using the is_valid_chapter function
    is_valid = is_valid_chapter(tanakh_division_name, book_name, chapter_choice, end_verse_choice)

    if is_valid:
        print(f"Chapter {chapter_choice} and Verse {end_verse_choice if end_verse_choice else ''} are valid!")
        return chapter_choice, start_verse_choice, end_verse_choice  # Return the valid chapter and verse values
    else:
        print(f"Invalid chapter or verse choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None  # Return None if invalid

def prompt_user_for_book(data):
    print("Please choose a section:")
    for key, value in data['sections'].items():
        print(f"{key}. {value}")
    
    tanakh_divisions = input("Enter the number corresponding to your choice: ")

    if tanakh_divisions == "1":
        tanakh_division_name = "Torah books"
    elif tanakh_divisions == "2":
        print(f"\n{tanakh_divisions}: Not yet coded - TODO")
        tanakh_division_name = "Prophets books"
        return None, None, None
    elif tanakh_divisions == "3":
        print(f"\n{tanakh_divisions}: Not yet coded - TODO")
        tanakh_division_name = "Scriptures books"
        return None, None, None
    else:
        print(f"Invalid choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None
    
    print(f"\nYou selected: {data['sections'][tanakh_divisions]}")
    print("\nPlease choose a book:")

    books = data[tanakh_division_name]
    for key, value in books.items():
        print(f"{key}. {value}")

    book_choice = input("Enter the number corresponding to your choice: ")

    if book_choice in books:
        book_name = books[book_choice]
        print(f"\nYou selected: {book_name}")
        return tanakh_division_name, book_choice, book_name
    else:
        print(f"Invalid choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return None, None, None

def is_valid_chapter(tanakh_division_name, book_choice, chapter_choice, verse_choice=None):
    if tanakh_division_name == TORAH_BOOKS:
        file_name = PENTATEUCH_FILE_ENG
    elif tanakh_division_name == PROPHETS_BOOKS:
        file_name = PROPHETS_FILE
    elif tanakh_division_name == SCRIPTURES_BOOKS:
        file_name = SCRIPTURES_FILE
    else:
        print(f"Invalid Tanakh division. line: {inspect.currentframe().f_lineno}: Exiting...")
        return False

    data = load_data(file_name)
    if data is None:
        return False

    if book_choice in data['books']:
        book_data = data['books'][book_choice]
    else:
        print(f"Invalid book choice: {book_choice}, read file: {file_name}. line: {inspect.currentframe().f_lineno}: Exiting...")
        return False

    if chapter_choice in book_data['chapters']:
        total_verses = book_data['chapters'][chapter_choice]
        
        if verse_choice is not None:
            if 1 <= verse_choice <= total_verses:
                return True
            else:
                print(f"Invalid verse choice. Chapter {chapter_choice} has {total_verses} verses.")
                return False
        else:
            return True
    else:
        print(f"Invalid chapter choice. line: {inspect.currentframe().f_lineno}: Exiting...")
        return False

def select_option_by_text(driver, select_name, option_text):
    """
    Select an option in a <select> dropdown by its visible text.

    :param driver: Selenium WebDriver instance.
    :param select_name: The 'name' attribute of the <select> element.
    :param option_text: The visible text of the option to select.
    """
    try:
        # Wait for the <select> element to be present
        select_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.NAME, select_name))
        )
        
        # Scroll into view
        driver.execute_script("arguments[0].scrollIntoView();", select_element)
        
        # Create a Select object and select the option by visible text
        select_object = Select(select_element)
        select_object.select_by_visible_text(option_text)
        
        print(f"Option '{option_text}' selected successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")

##################################################################################
# Generic function to select an option from a dropdown
##################################################################################
def select_option(driver, dropdown_name, option_text):
    try:
        dropdown = Select(driver.find_element(By.NAME, dropdown_name))
        dropdown.select_by_visible_text(option_text)
        print(f"Option '{option_text}' selected from dropdown '{dropdown_name}'.")
    except Exception as e:
        print(f"An error occurred while selecting '{option_text}' from '{dropdown_name}': {e}")

# Function to click a specific link
def click_link(driver, link_text):
    try:
        link = driver.find_element(By.LINK_TEXT, link_text)
        link.click()
        print(f"Link with text '{link_text}' clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the link '{link_text}': {e}")

##################################################################################
# Function to click a submit button
##################################################################################
def click_submit_button(driver):
    try:
        submit_button = driver.find_element(By.XPATH, "//input[@type='submit' and @value='GO']")
        submit_button.click()
        print("Submit button clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the submit button: {e}")

# Function to get the current URL
def get_current_url(driver):
    try:
        current_url = driver.current_url
        print(f"Current URL: {current_url}")
        return current_url
    except Exception as e:
        print(f"An error occurred while retrieving the current URL: {e}")
        return None

##################################################################################
# Grabs all verses and their text from the web page.
##################################################################################
def grab_verses(driver):
    """
    Grabs all verses and their text from the web page.

    Parameters:
        driver: WebDriver instance
    Returns:
        List of tuples containing verse number and text
    """
    try:
        # Wait for the page to load completely
        WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.TAG_NAME, "b")))

        # Get the page source after it has fully loaded
        page_source = driver.page_source

        # Parse the page with BeautifulSoup
        soup = BeautifulSoup(page_source, "html.parser")

        # Find all <b> tags that represent the verses
        verses = soup.find_all('b')

        verse_texts = []

        # Loop through all the <b> tags to extract the verse numbers and text
        for verse in verses:
            verse_number = verse.get_text(strip=True)  # Get verse number (e.g., "Verse 1")

            # Try to find the text either directly following the <b> tag or within a <p> tag
            verse_text = ""
            if verse.next_sibling:  # Check if there's text directly after the <b> tag
                verse_text = verse.next_sibling.strip()
            else:
                next_p = verse.find_next('p')  # Check for text in the next <p> tag
                if next_p:
                    verse_text = next_p.get_text(strip=True)

            if verse_number and verse_text:
                verse_texts.append((verse_number, verse_text))  # Store as a tuple

        return verse_texts

    except Exception as e:
        print(f"An error occurred while grabbing verses: {e}")
        return []

##################################################################################
# Function to save verses to a Word document
##################################################################################
def save_to_word(verses, filename, book_name, chapter_number, file_path="."):
    """
    Save a list of verses to a Word document.

    Parameters:
        verses (list of tuples): Each tuple contains a verse number and verse text.
        filename (str): The base name of the Word file to save the verses.
        book_name (str): Name of the book (e.g., "Genesis").
        chapter_choice (int): Chapter number.
        start_verse_choice (int): Starting verse number.
        end_verse_choice (int): Ending verse number.
        file_path (str): Directory path to save the file. Default is the current directory.
    """
    # Define the file path dynamically
    os.makedirs(file_path, exist_ok=True)
    save_path = os.path.join(
        file_path,
        f"{filename}.docx"
    )

    # Delete the file if it exists
    if os.path.exists(save_path):
        print(f"File exists, deleting: {save_path}")
        os.remove(save_path)

    # Create a new Word Document
    doc = Document()

    # Add a title to the document
    doc.add_heading(f'{filename} - Chapter {chapter_number}', 0)

    # Add each verse to the document
    for verse_number, verse_text in verses:
        doc.add_paragraph(f"{verse_number}: {verse_text}")

    # Save the document
    doc.save(save_path)
    print(f"Saved Hebrew-friendly Word document: {save_path}")


##################################################################################
# Function to open a website in Google Chrome on macOS
##################################################################################
def main_open_website_with_chrome(website_url):
    try:
        # Path to the Google Chrome executable on macOS
        chrome_path = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        subprocess.Popen([chrome_path, website_url])  # Launch Chrome without tying it to the Python script
        print(f"Website {website_url} opened in Google Chrome successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")

##################################################################################
# Get any chapter from any book. 
##################################################################################
def main_tanakh_ch():
    # Get user inputs
    inputs = get_tanakh_scraper_inputs()
    if not inputs:
        return

    tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice = inputs

    chapter_number = chapter_choice.zfill(2)
    get_Tanakh_and_verses(chapter_number, book_name)  # Pass the chapter number to the function

##################################################################################
# Click links and scrape to get text to a word document based on parameters. 
##################################################################################
def get_Tanakh_and_verses(chapter_number, book_name):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://www.mnemotrix.com/texis/vtx/chumash")  # Replace with your desired URL

    try:
        # Step 1: Select options to get to the next page
        select_option(driver, "bookq", book_name)         # Select "Genesis" from the "bookq" dropdown
        select_option(driver, "chapterq", f"Chapter {chapter_number}")  # Use chapter_number as a string
        time.sleep(3)
        click_submit_button(driver)  # Click the submit button

        # Step 2 click the link to the second page
        # Locate and click the first link with the specified text
        time.sleep(2)
        partial_text = get_partial_text(book_name)
        links = driver.find_elements(By.PARTIAL_LINK_TEXT, partial_text)
        if links:  # Check if any links were found
            links[0].click()  # Click the first matching link
        else:
            print(f"----> No matching link found for {book_name} ch: {chapter_number}")
        time.sleep(2)
        # Step 3: Retrieve and print the final URL
        final_url = get_current_url(driver)

        # Step 4: Grab all verses from the page using the final URL
        driver.get(final_url)  # Navigate to the final URL
        verses = grab_verses(driver)

        # Step 5: Save the verses to a Word document with the specified filename format
        filename = f"{book_name}_{chapter_number}.docx"
        folder_path=load_tanakh_path(ENG_DOCX_FOLDER)
        folder_path = os.path.join(folder_path, book_name)
        save_to_word(verses, filename, book_name, chapter_number, file_path=folder_path)

    finally:
        # Step Last: Wait and Quit
        time.sleep(3)  # Wait for 5 seconds to observe the result
        driver.quit()

def get_partial_text(book_name):
    book_mapping = {
        "Genesis": "Bereishis/Genesis, Chapter",
        "Exodus": "Shemos/Exodus, Chapter",
        "Leviticus": "Vayikro/Leviticus, Chapter",
        "Numbers": "Bamidbar/Numbers, Chapter",
        "Deuteronomy": "Devarim/Deuteronomy, Chapter",
    }
    return book_mapping.get(book_name, "Error: Invalid book input")

def main_torah_book_eng():
    # Map books to their chapter ranges
    books = {
        "Genesis": 51,
        "Exodus": 40,
        "Leviticus": 27,
        "Numbers": 36,
        "Deuteronomy": 34,
    }

    # Prompt the user for input
    print("Choose a book of the Torah to process:")
    for i, book in enumerate(books.keys(), 1):
        print(f"{i}. {book}")
    
    try:
        choice = int(input("Enter the number corresponding to your choice: "))
        if choice < 1 or choice > len(books):
            raise ValueError("Invalid choice. Please choose a number from the list.")

        # Get the selected book and its chapter range
        selected_book = list(books.keys())[choice - 1]
        chapter_range = books[selected_book]

        print(f"Processing {selected_book}...")
        for chapter_number in range(1, chapter_range + 1):
            time.sleep(3)
            # Convert the chapter number to a two-digit string if necessary
            chapter_number_str = str(chapter_number).zfill(2)
            # Get the verses for the chapter
            get_Tanakh_and_verses(chapter_number_str, selected_book)
        
        print(f"Finished processing {selected_book}.")

    except ValueError as e:
        print(f"Error: {e}. Please restart and enter a valid number.")

def process_all_parashot_main(file_path="data/torah_parashot_eng.json"):
    """
    Processes all the parashot in the torah_parashot.json file by passing their names to
    the get_tanakh_range_from_json_main function.

    :param file_path: Path to the torah_parashot.json file.
    """
    try:
        # Load the parasha data from the JSON file
        with open(file_path, 'r', encoding='utf-8') as file:
            parasha_data = json.load(file)
        
        # Loop through all the "Name" fields and process each parasha
        for parasha in parasha_data.get("Parashot", []):
            parasha_name = parasha.get("Name")
            if parasha_name:
                print(f"Processing parasha: {parasha_name}")
                time.sleep(20)
                get_tanakh_range_from_json_main(parasha_name, file_path)
            else:
                print("Skipping a parasha with missing 'Name' field.")
    
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: Failed to decode JSON from '{file_path}'.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def get_tanakh_range_from_json_main(parasha_name, file_path="data/torah_parashot.json"):
    """
    Gets the Tanakh range based on the specified parasha name from torah_parashot.json.

    :param parasha_name: Name of the Torah portion (e.g., "Bereshit", "Noach").
    :param file_path: Path to the torah_parashot.json file.
    """
    try:
        # Load the parasha data from the JSON file
        with open(file_path, 'r', encoding='utf-8') as file:
            parasha_data = json.load(file)
        
        # Find the specified parasha
        parasha = next((p for p in parasha_data["Parashot"] if p["Name"] == parasha_name), None)
        
        if not parasha:
            print(f"Error: Parasha '{parasha_name}' not found in the data.")
            return
        
        # Extract range information
        tanakh_division_name = parasha["Tanakh Section"]
        book_name = parasha["Book"]
        start_chapter = parasha["Start"]["Chapter"]
        start_verse = parasha["Start"]["Verse"]
        end_chapter = parasha["End"]["Chapter"]
        end_verse = parasha["End"]["Verse"]

        start_chapter = str(start_chapter).zfill(2)
        end_chapter = str(end_chapter).zfill(2)
        
        # Traverse and scrape using the extracted range
        folder_path=load_tanakh_path(ENG_DOCX_FOLDER)
        folder_path = os.path.join(folder_path, book_name)
        traverse_tanakh_scraper(
            parasha_name=parasha_name,
            book_name=book_name,
            chapter_choice=start_chapter,
            end_chapter_choice=end_chapter,
            start_verse_choice=start_verse,
            end_verse_choice=end_verse,
            file_path=folder_path
        )
    
    except FileNotFoundError:
        print(f"Error: The file '{file_path}' was not found.")
    except json.JSONDecodeError:
        print(f"Error: Failed to decode JSON from '{file_path}'.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def traverse_tanakh_scraper(parasha_name, book_name=None, chapter_choice=None, end_chapter_choice=None, start_verse_choice=1, end_verse_choice=None, file_path=load_tanakh_path(ENG_DOCX_FOLDER)):
    """
    Traverses the Tanakh JSON structure and performs scraping for specified sections.
    """
    DEBUG = True  # Toggle for debug print statements

    try:
        # Load the JSON file
        with open(os.path.join("data", "Pentateuch.json"), 'r', encoding='utf-8') as file:
            tanakh_data = json.load(file)

        books = tanakh_data.get("books", {})
        
        # Traverse books in the division
        for current_book_name, book_data in books.items():
            if book_name and current_book_name != book_name:
                continue  # Skip if not the specified book
            
            chapters = book_data.get("chapters", {})
            
            for current_chapter, verse_count in chapters.items():
                current_chapter = int(current_chapter)  # Convert chapter to integer for comparison
                
                if chapter_choice and current_chapter < int(chapter_choice):
                    continue  # Skip chapters before the starting chapter
                
                if end_chapter_choice and current_chapter > int(end_chapter_choice):
                    break  # Stop processing beyond the ending chapter
                
                # Set the start and end verse ranges
                start_verse = start_verse_choice if current_chapter == chapter_choice else 1
                end_verse = end_verse_choice if current_chapter == end_chapter_choice else verse_count

                if DEBUG:
                    print(f"Processing {current_book_name}, Chapter {current_chapter}, Verses {start_verse}-{end_verse}")

                time.sleep(1) 
                # Perform scraping for the current range
                get_Tanakh_Parashot(parasha_name, current_chapter, book_name)
                
    except FileNotFoundError:
        print(f"Error: The file 'Pentateuch.json' was not found in the 'data' folder.")
    except json.JSONDecodeError:
        print("Error: Failed to decode JSON from 'Pentateuch.json'.")
    except Exception as e:
        if DEBUG:
            print(f"An unexpected error occurred: {e}")

def get_Tanakh_Parashot(parasha_name, chapter_number, book_name):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://www.mnemotrix.com/texis/vtx/chumash")  # Replace with your desired URL

    print(f"Parasha name: {parasha_name}")

    try:
        # Step 1: Select options to get to the next page
        select_option_by_text(driver, "itemq", parasha_name)
        select_option(driver, "bookq", book_name)         # Select "Genesis" from the "bookq" dropdown
        select_option(driver, "chapterq", f"Chapter {chapter_number}")  # Use chapter_number as a string
        time.sleep(3)
        click_submit_button(driver)  # Click the submit button

        # Step 2 click the link to the second page
        # Locate and click the first link with the specified text
        time.sleep(2)
        partial_text = get_partial_text(book_name)
        links = driver.find_elements(By.PARTIAL_LINK_TEXT, partial_text)
        if links:  # Check if any links were found
            links[0].click()  # Click the first matching link
        else:
            print(f"----> No matching link found for {book_name} ch: {chapter_number}")
        time.sleep(2)
        # Step 3: Retrieve and print the final URL
        final_url = get_current_url(driver)

        # Step 4: Grab all verses from the page using the final URL
        driver.get(final_url)  # Navigate to the final URL
        verses = grab_verses(driver)

        # Step 5: Save the verses to a Word document with the specified filename format
        filename = f"{book_name}_{chapter_number}.docx"
        folder_path=load_tanakh_path(ENG_DOCX_FOLDER)
        folder_path = os.path.join(folder_path, parasha_name)
        save_to_word(verses, filename, book_name, chapter_number, file_path=folder_path)

    finally:
        # Step Last: Wait and Quit
        time.sleep(3)  # Wait for 5 seconds to observe the result
        driver.quit()

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
def prompt_user_choice():
    # Ask the user to choose between the options
    print("Choose an option:")
    print("1. Open english Torah Site")
    print("2. Get any chapter in the Torah")
    print("3. Get any book of the Torah")
    print("4. Get any parasha of the Torah")
    print("5. Get all parashot of the Torah")
    print("6. Print all parashot of the Torah")

    choice = input("Please enter a number: 1 through 3.: ").strip()
    file_path = load_data(PARASHOT_LIST_ENG_FILE, return_path_only=True)

    if choice == "1":
        # Call the function to get all Genesis chapters from 1 to 50
        eng_website_url = "http://www.mnemotrix.com/texis/vtx/chumash"
        main_open_website_with_chrome(eng_website_url)
    elif choice == "2":
        # Call the function to get a specific Genesis chapter
        main_tanakh_ch()
    elif choice == "3":
        main_torah_book_eng()
    elif choice == "4":
    # Get user inputs
        inputs = get_tanakh_scraper_inputs()
        if not inputs:
            return
        tanakh_division_name, book_name, chapter_choice, start_verse_choice, end_verse_choice = inputs

        #print_parashah_info_main(file_name)
        chapter_number = chapter_choice.zfill(2)
        parasha_name = "Breishit"
        get_Tanakh_Parashot(parasha_name, chapter_number, book_name)
    elif choice == "5":
        process_all_parashot_main(file_path)
    if choice == "6":
        # Call the Parashah info printing function
        print_parashah_info_main(file_path)
    else:
        print("Invalid choice. Please enter a number: 1 through 4.")

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
if __name__ == "__main__":
    # Ask the user to choose between the options - retrieve data 
    prompt_user_choice()