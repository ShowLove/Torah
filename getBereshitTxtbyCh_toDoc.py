from selenium import webdriver                                      # For automating and controlling the web browser
from selenium.webdriver.common.by import By                         # For locating elements on the web page
from selenium.webdriver.chrome.service import Service               # For initializing and configuring the ChromeDriver
from webdriver_manager.chrome import ChromeDriverManager            # Manages downloading and setting up ChromeDriver
from selenium.webdriver.support.ui import Select                    # For interacting with drop-down menus (select elements)
import time                                                         # For pausing the execution of the program (e.g., sleep or wait)
import subprocess                                                   # For running system commands and interacting with the system shell
from bs4 import BeautifulSoup                                       # For parsing and navigating HTML or XML content
from selenium.webdriver.support.ui import WebDriverWait             # For waiting for elements to appear on the page
from selenium.webdriver.support import expected_conditions as EC    # For defining the expected conditions for elements
from docx import Document                                           # For creating and modifying Word documents
import os                                                           # For file and directory operations (e.g., working with paths, creating folders)
import shutil                                                       # For file operations (e.g., moving, copying, and deleting files)
                                                                    ################################################################################################

##################################################################################
# Generic function to select an option from a dropdown
##################################################################################
def select_option(driver, dropdown_name, option_text):
    try:
        dropdown = Select(driver.find_element(By.NAME, dropdown_name))
        dropdown.select_by_visible_text(option_text)
        print(f"Option '{option_text}' selected from dropdown '{dropdown_name}'.")
    except Exception as e:
        print(f"An error occurred while selecting '{option_text}' from '{dropdown_name}': {e}")

# Function to click a specific link
def click_link(driver, link_text):
    try:
        link = driver.find_element(By.LINK_TEXT, link_text)
        link.click()
        print(f"Link with text '{link_text}' clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the link '{link_text}': {e}")

##################################################################################
# Function to click a submit button
##################################################################################
def click_submit_button(driver):
    try:
        submit_button = driver.find_element(By.XPATH, "//input[@type='submit' and @value='GO']")
        submit_button.click()
        print("Submit button clicked.")
    except Exception as e:
        print(f"An error occurred while clicking the submit button: {e}")

# Function to get the current URL
def get_current_url(driver):
    try:
        current_url = driver.current_url
        print(f"Current URL: {current_url}")
        return current_url
    except Exception as e:
        print(f"An error occurred while retrieving the current URL: {e}")
        return None

##################################################################################
# Grabs all verses and their text from the web page.
##################################################################################
def grab_verses(driver):
    """
    Grabs all verses and their text from the web page.

    Parameters:
        driver: WebDriver instance
    Returns:
        List of tuples containing verse number and text
    """
    try:
        # Wait for the page to load completely
        WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.TAG_NAME, "b")))

        # Get the page source after it has fully loaded
        page_source = driver.page_source

        # Parse the page with BeautifulSoup
        soup = BeautifulSoup(page_source, "html.parser")

        # Find all <b> tags that represent the verses
        verses = soup.find_all('b')

        verse_texts = []

        # Loop through all the <b> tags to extract the verse numbers and text
        for verse in verses:
            verse_number = verse.get_text(strip=True)  # Get verse number (e.g., "Verse 1")

            # Try to find the text either directly following the <b> tag or within a <p> tag
            verse_text = ""
            if verse.next_sibling:  # Check if there's text directly after the <b> tag
                verse_text = verse.next_sibling.strip()
            else:
                next_p = verse.find_next('p')  # Check for text in the next <p> tag
                if next_p:
                    verse_text = next_p.get_text(strip=True)

            if verse_number and verse_text:
                verse_texts.append((verse_number, verse_text))  # Store as a tuple

        return verse_texts

    except Exception as e:
        print(f"An error occurred while grabbing verses: {e}")
        return []

##################################################################################
# Function to save verses to a Word document
##################################################################################
def save_to_word(verses, filename):
    """
    Save a list of verses to a Word document.

    Parameters:
        verses (list of tuples): Each tuple contains a verse number and verse text.
        filename (str): The name of the Word file to save the verses.
    """
    # Create a new Document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Genesis - Verses', 0)

    # Add each verse to the document
    for verse_number, verse_text in verses:
        doc.add_paragraph(f"{verse_number}: {verse_text}")

    # Save the document
    doc.save(filename)
    print(f"Verses have been saved to {filename}")

##################################################################################
# Extract the full link string from book and chapter information
##################################################################################
def extract_full_string(html_content):
    # Parse the HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Find all links in the HTML content
    links = soup.find_all('a', href=True)
    
    # Loop through all links to find the partial string "Bereishis/Genesis, Chapter 01"
    for link in links:
        # Check if the link text contains the partial string
        if 'Bereishis/Genesis, Chapter' in link.get_text():
            # Prepend "Bereishis: " to the found link text to get the full string
            full_string = link.get_text().strip()
            return full_string
    
    return None  # Return None if the string is not found

##################################################################################
# Main process to get Genesis and feed the URL into the verse-grabbing function
##################################################################################
def get_Genesis_and_verses(chapter_number):
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://www.mnemotrix.com/texis/vtx/chumash")  # Replace with your desired URL

    try:
        # Step 1: Select options to get to the next page
        select_option(driver, "bookq", "Genesis")         # Select "Genesis" from the "bookq" dropdown
        select_option(driver, "chapterq", f"Chapter {chapter_number}")  # Use chapter_number as a string
        click_submit_button(driver)  # Click the submit button

        # Step 2: Click the link on the second page
        html_content = driver.page_source
        link_text = extract_full_string(html_content)
        #link_text = f"Bereishis/Genesis, Chapter {chapter_number}"  # Use chapter_number in the link text
        click_link(driver, link_text)

        # Step 3: Retrieve and print the final URL
        final_url = get_current_url(driver)

        # Step 4: Grab all verses from the page using the final URL
        driver.get(final_url)  # Navigate to the final URL
        verses = grab_verses(driver)

        # Step 5: Save the verses to a Word document with the specified filename format
        filename = f"gen_{chapter_number}.docx"
        save_to_word(verses, filename)

    finally:
        # Step Last: Wait and Quit
        time.sleep(5)  # Wait for 5 seconds to observe the result
        driver.quit()

##################################################################################
# Encapsulated main function
##################################################################################
def main_get_gen_ch():
    # Prompt the user for the chapter number between 1 and 50
    chapter_number = input("Enter the chapter number (1-50): ").strip()

    # Validate the chapter number input
    if chapter_number.isdigit() and 1 <= int(chapter_number) <= 50:
        # Convert to two-digit string if necessary
        chapter_number = chapter_number.zfill(2)
        get_Genesis_and_verses(chapter_number)  # Pass the chapter number to the function
    else:
        print("Invalid chapter number. Please enter a number between 1 and 50.")

def main_get_gen():
    # Iterate through chapters 1 to 50
    for chapter_number in range(1, 51):
        # Convert the chapter number to a two-digit string if necessary
        chapter_number_str = str(chapter_number).zfill(2)
        # Get the verses for the chapter
        get_Genesis_and_verses(chapter_number_str)

##################################################################################
# Function to open a website in Google Chrome on macOS
##################################################################################
def main_open_website_with_chrome(website_url):
    try:
        # Path to the Google Chrome executable on macOS
        chrome_path = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        subprocess.Popen([chrome_path, website_url])  # Launch Chrome without tying it to the Python script
        print(f"Website {website_url} opened in Google Chrome successfully.")
    except Exception as e:
        print(f"An error occurred: {e}")

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
def prompt_user_choice():
    # Ask the user to choose between the options
    print("Choose an option:")
    print("1. Get Genesis Chapter for a specific number")
    print("2. Get Genesis Chapters 1-50")
    print("3. Open english Torah Site")

    choice = input("Please enter a number: 1 through 3.: ").strip()

    if choice == "1":
        # Call the function to get a specific Genesis chapter
        main_get_gen_ch()
    elif choice == "2":
        # Call the function to get all Genesis chapters from 1 to 50
        main_get_gen()
    elif choice == "3":
        # Call the function to get all Genesis chapters from 1 to 50
        eng_website_url = "http://www.mnemotrix.com/texis/vtx/chumash"
        main_open_website_with_chrome(eng_website_url)
    else:
        print("Invalid choice. Please enter a number: 1 through 3.")
        prompt_user_choice()  # Recurse until a valid choice is made

##################################################################################
# Move all the current word files in current directory to a specific folder. 
##################################################################################
def move_word_files_to_folder(destination_folder):
    # Get the current directory
    current_directory = os.getcwd()

    # Ensure the destination folder exists, if not create it
    if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)

    # Loop through all files in the current directory
    for file_name in os.listdir(current_directory):
        # Check if the file is a Word document (.docx)
        if file_name.endswith('.docx'):
            # Construct the full file path
            source_file = os.path.join(current_directory, file_name)
            destination_file = os.path.join(destination_folder, file_name)

            # Move the file (replace if it already exists)
            try:
                shutil.move(source_file, destination_file)
                print(f"Moved: {file_name}")
            except Exception as e:
                print(f"Error moving {file_name}: {e}")

##################################################################################
# Call prompt_user_choice in the main entry point
##################################################################################
if __name__ == "__main__":
    # Ask the user to choose between the options - retrieve data 
    prompt_user_choice()

    # put the data in correct directory if needed
    subfolder_name = 'bereshit_eng_files'  # Replace with your folder path
    current_directory = os.getcwd()
    destination_folder = os.path.join(current_directory, subfolder_name)
    move_word_files_to_folder(destination_folder)