from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import re
import os
from docx.shared import RGBColor

# File Paths
TANAKH_DOCX_FOLDER = "tanakh_docs"
ENG_DOCX_FOLDER = "eng_docs"
HEB_DOCX_FOLDER = "hebrew_docs"
OUTPUT_DOCX_FOLDER = "output_docs"

# Configurable constants for formatting
DOCX_HEBREW_FONT = "Frank Ruehl"  # Use Frank Ruehl for Hebrew text on Word
DOCX_ENGLISH_FONT = "Times New Roman"  # Use Times New Roman for English text
FONT_SIZE_HEB = 16  # Font size in points
FONT_SIZE_ENG = 12  # Font size in points
MARGIN_SIZE = Pt(12)  # Margin size in points

def load_tanakh_path(folder_name):
    file_path = os.path.join(TANAKH_DOCX_FOLDER, folder_name)
    return file_path

def navigate_folders_from_cwd(start_path):
    """
    Dynamically navigate folders starting from the current working directory and return the final path.

    Returns:
    - str: The final folder path.
    """
    current_path = start_path

    while True:
        print(f"\nCurrent Path: {current_path}")
        subfolders = [f for f in os.listdir(current_path) if os.path.isdir(os.path.join(current_path, f))]

        if not subfolders:
            print("No subfolders available in the current directory.")
            break

        # Display subfolders with options
        print("Subfolders:")
        for i, folder in enumerate(subfolders, 1):
            print(f"{i}. {folder}")
        print("0. Select this folder")
        print("-1. Go up one level")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            # Select the current folder
            return current_path
        elif choice == -1:
            # Go up one level
            parent_path = os.path.dirname(current_path)
            if current_path == parent_path:  # At the root
                print("You are already at the root directory.")
            else:
                current_path = parent_path
        elif 1 <= choice <= len(subfolders):
            # Navigate into the selected subfolder
            current_path = os.path.join(current_path, subfolders[choice - 1])
        else:
            print("Invalid choice. Please try again.")

    return current_path

def pick_filename_from_folder(folder_path):
    """
    Dynamically pick a filename from the specified folder path.

    Parameters:
    - folder_path (str): The path to the folder.

    Returns:
    - str: The selected filename, or None if no valid file is selected.
    """
    if not os.path.isdir(folder_path):
        print("Invalid folder path. Please provide a valid directory.")
        return None

    while True:
        print(f"\nCurrent Folder: {folder_path}")
        files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

        if not files:
            print("No files available in this folder.")
            return None

        # Display files with options
        print("Files:")
        for i, file in enumerate(files, 1):
            print(f"{i}. {file}")
        print("0. Exit without selecting a file")

        # Get user input
        try:
            choice = int(input("Enter your choice: ").strip())
        except ValueError:
            print("Invalid input. Please enter a number.")
            continue

        if choice == 0:
            print("No file selected.")
            return None
        elif 1 <= choice <= len(files):
            # Return the selected file
            selected_file = files[choice - 1]
            return os.path.join(folder_path, selected_file)
        else:
            print("Invalid choice. Please try again.")

def weave_torah_files(parasha_name, hebrew_file_path, english_file_path, output_file_path):
    """
    Combine Hebrew and English texts from Word documents into a single Word file.
    Preserve original formatting of Hebrew and English documents.

    Args:
        hebrew_file_path (str): Path to the Hebrew Word document.
        english_file_path (str): Path to the English Word document.
        output_file_path (str): Path for the combined Word document output.
    """
    # Load Hebrew and English Word files
    hebrew_doc = Document(hebrew_file_path)
    english_doc = Document(english_file_path)

    # Prepare the output Word document
    output_doc = Document()

    # Set narrow margins
    section = output_doc.sections[0]
    section.left_margin = Inches(0.5)   # 0.5 inches left margin
    section.right_margin = Inches(0.5)  # 0.5 inches right margin
    section.top_margin = Inches(0.5)    # 0.5 inches top margin
    section.bottom_margin = Inches(0.5) # 0.5 inches bottom margin

    hebrew_index = 0
    english_index = 0

    # Process paragraphs and align Hebrew and English text
    while hebrew_index < len(hebrew_doc.paragraphs):
        # Add the Hebrew paragraph with formatting
        hebrew_paragraph = hebrew_doc.paragraphs[hebrew_index]
        output_paragraph = output_doc.add_paragraph()
        
        # Set the Hebrew font and size
        for run in hebrew_paragraph.runs:
            run_new = output_paragraph.add_run(run.text)
            run_new.bold = run.bold
            run_new.italic = run.italic
            run_new.underline = run.underline
            run_new.font.name = DOCX_HEBREW_FONT
            run_new.font.size = Pt(FONT_SIZE_HEB)

        # Align Hebrew text to the right (Right-to-left alignment)
        output_paragraph.alignment = 1  # Center-aligned (Note: 2 corresponds to right in python-docx)
        # Ensure the text direction is right-to-left for Hebrew
        output_paragraph.paragraph_format.bidi = True  # This forces the direction to be right-to-left for Hebrew

        hebrew_index += 1

        # Add the corresponding English paragraph with formatting
        if english_index < len(english_doc.paragraphs):
            english_paragraph = english_doc.paragraphs[english_index]
            output_paragraph = output_doc.add_paragraph()

            # Set the English font and size
            for run in english_paragraph.runs:
                run_new = output_paragraph.add_run(run.text)
                run_new.bold = run.bold
                run_new.italic = run.italic
                run_new.underline = run.underline
                run_new.font.name = DOCX_ENGLISH_FONT
                run_new.font.size = Pt(FONT_SIZE_ENG)

            # Align English text to the left (Left-to-right alignment)
            output_paragraph.alignment = 0  # Left-aligned (0 corresponds to left in python-docx)
            # Ensure the text direction is left-to-right for English
            output_paragraph.paragraph_format.bidi = False  # This forces the direction to be left-to-right for English

            english_index += 1

    # Check if Hebrew paragraphs exist and prepare the first line for the file name
    if hebrew_doc.paragraphs:
        first_line_hebrew = hebrew_doc.paragraphs[0].text.strip()
        # Clean the Hebrew text to use it safely in the filename (replace or remove invalid characters)
        safe_first_line_hebrew = "".join([c if c.isalnum() or c.isspace() else "_" for c in first_line_hebrew])
    else:
        safe_first_line_hebrew = "combined"

    # Ensure the output path exists
    if not os.path.exists(output_file_path):
        os.makedirs(output_file_path)

    # Define the output file path using the cleaned Hebrew text
    docx_name = f"{parasha_name}_{safe_first_line_hebrew}.docx"
    docx_name = clean_hebrew_filename(docx_name)
    final_path_result = os.path.join(output_file_path, docx_name)

    # Remove the existing file if it already exists
    if os.path.exists(final_path_result):
        os.remove(final_path_result)
        print(f"Existing file found and removed: {final_path_result}")

    # Save the combined document
    output_doc.save(final_path_result)
    print(f"Combined document saved as: {final_path_result}")

def clean_hebrew_filename(filename):
    """
    Clean a Hebrew filename string to a consistent format.
    
    Args:
        filename (str): The original filename string.
        
    Returns:
        str: The cleaned filename string.
    """
    # Replace spaces with underscores
    cleaned_string = re.sub(r"\s+", "_", filename)
    # Replace multiple underscores with a single underscore
    cleaned_string = re.sub(r"_{2,}", "_", cleaned_string)
    # Ensure proper format for "Chapter"
    cleaned_string = re.sub(r"_Chapter_", " Chapter_", cleaned_string)
    # Ensure proper format for "Verses"
    cleaned_string = re.sub(r"_Verses_", " Verses_", cleaned_string)
    # Remove extra underscores after "Chapter" and "Verses"
    cleaned_string = re.sub(r"(?<=Chapter)_+", "_", cleaned_string)
    cleaned_string = re.sub(r"(?<=Verses)_+", "_", cleaned_string)
    # Remove trailing underscores
    cleaned_string = re.sub(r"_+$", "", cleaned_string)
    # Handle dangling underscores before ".docx"
    cleaned_string = re.sub(r"_\.docx$", ".docx", cleaned_string)
    
    return cleaned_string

def add_notes_to_verses(file_path):
    """
    Adds a notes section to each verse in a .docx file while preserving formatting 
    and replaces the original file. The output file has narrow margins.

    Args:
        file_path (str): Path to the .docx file to process.
    """
    # Load the document
    doc = Document(file_path)
    
    # Extract the header (assumes it's the first paragraph with the chapter name)
    header_text = None
    for para in doc.paragraphs:
        if "Chapter" in para.text:
            header_text = para.text.strip()
            break
    
    if not header_text:
        raise ValueError("No header containing 'Chapter' found in the document.")
    
    # Create a new document to rewrite the content
    new_doc = Document()

    # Set narrow margins for the new document
    section = new_doc.sections[0]
    section.left_margin = Inches(0.5)   # 0.5 inches left margin
    section.right_margin = Inches(0.5)  # 0.5 inches right margin
    section.top_margin = Inches(0.5)    # 0.5 inches top margin
    section.bottom_margin = Inches(0.5) # 0.5 inches bottom margin

    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        # Copy the paragraph while preserving formatting
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.alignment = para.alignment  # Preserve alignment
        for run in para.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
        
        # Check if the paragraph is a verse
        if para_text.startswith("Verse"):
            # Extract the verse number using regex
            match = re.match(r"Verse (\d+)", para_text)
            if match:
                verse_number = match.group(1)
                # Add the notes paragraph for the verse
                notes_paragraph = new_doc.add_paragraph()
                notes = f"[notes]( {header_text} Verse {verse_number} )[end_notes]"
                notes_paragraph.alignment = 0  # Align notes to the left (customize if needed)
                notes_run = notes_paragraph.add_run(notes)
                # Optional: Style the notes text (customize if needed)
                notes_run.italic = True
                notes_run.font.color.rgb = RGBColor(211, 211, 211)  # Set the notes text to light grey
    
    # Delete the original file if it exists
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Existing file deleted: {file_path}")
    
    # Save the reformatted content back to the original file
    new_doc.save(file_path)
    print(f"Formatted document saved as: {file_path}")

def remove_second_colon(para_text, paragraph):
    # This regex will match ": (כח):" and remove the second colon at the end
    para_text = re.sub(r"(:\((\D+)\)‪):", r" (\2)‪:", para_text)
    
    # Now we return the modified paragraph with the second colon removed
    new_paragraph = paragraph
    new_paragraph.clear()  # Clear previous runs
    new_paragraph.add_run(para_text)  # Add the modified text without the second colon
    
    return new_paragraph

def format_docx_file(file_path):
    """
    Formats the content of a .docx file, applies consistent formatting while preserving
    the content as-is, and replaces the original file.

    Args:
        file_path (str): Path to the .docx file to process.
    """
    # Load the document
    doc = Document(file_path)
    
    # Create a new document to rewrite the content
    new_doc = Document()

    # Set narrow margins for the new document
    section = new_doc.sections[0]
    section.left_margin = Inches(0.5)   # 0.5 inches left margin
    section.right_margin = Inches(0.5)  # 0.5 inches right margin
    section.top_margin = Inches(0.5)    # 0.5 inches top margin
    section.bottom_margin = Inches(0.5) # 0.5 inches bottom margin

    # Regex to detect Hebrew characters
    hebrew_range = r"[\u0590-\u05FF]"

    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        # Create a new paragraph in the formatted document
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.alignment = para.alignment  # Preserve alignment
        
        # Loop through the runs in the original paragraph and reapply formatting
        for run in para.runs:
            new_run = new_paragraph.add_run(run.text)
            
            # Preserve original run formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb

            # Apply consistent font and size based on content
            if re.search(hebrew_range, run.text):  # If the text contains Hebrew characters
                new_run.font.name = DOCX_HEBREW_FONT
                new_run.font.size = Pt(FONT_SIZE_HEB)
            else:  # Non-Hebrew text
                new_run.font.name = DOCX_ENGLISH_FONT
                new_run.font.size = Pt(FONT_SIZE_ENG)

    # Delete the original file if it exists
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Existing file deleted: {file_path}")
    
    # Save the reformatted content back to the original file
    new_doc.save(file_path)
    print(f"Formatted document saved as: {file_path}")

if __name__ == "__main__":

    # Step 1: get the paths
    eng_folder_path = load_tanakh_path(ENG_DOCX_FOLDER)
    heb_folder_path = load_tanakh_path(HEB_DOCX_FOLDER)
    output_folder_path = load_tanakh_path(OUTPUT_DOCX_FOLDER)
    hebrew_file_path = "hebrew_text.docx"
    english_file_path = "english_text.docx"
    output_file_path = "combined_output.docx"

    final_eng_path = navigate_folders_from_cwd(eng_folder_path)
    print(f"\nEng - Final Path: {final_eng_path}")
    english_file = pick_filename_from_folder(final_eng_path)
    if english_file:
        print(f"\nSelected File: {english_file}")

    final_heb_path = navigate_folders_from_cwd(heb_folder_path)
    print(f"\nHeb - Final Path: {final_heb_path}")
    heb_file = pick_filename_from_folder(final_heb_path)
    parasha_name = os.path.basename(final_heb_path)
    if heb_file:
        print(f"\nSelected File: {heb_file}")
        print(f"\nParasha Name: {parasha_name}")

    # Step 2: weave heb_file and english_file
    weave_torah_files(parasha_name, heb_file, english_file, output_folder_path)

    # Step 3 do post processing on the weaved together output
    final_output_file = pick_filename_from_folder(output_folder_path)
    if final_output_file:
        print(f"\nSelected File: {final_output_file}")

    #add_notes_to_verses(final_output_file)
    doc = Document(final_output_file)

    # Iterate over paragraphs and remove second Heb verse number colon
    for para in doc.paragraphs:
        para_text = para.text.strip()
        new_paragraph = remove_second_colon(para_text, para)

    # Save the modified document
    doc.save(final_output_file)

    format_docx_file(final_output_file)

