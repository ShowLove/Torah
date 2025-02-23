import utils

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import re
import os
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def get_full_filename(partial_filename, search_dir):
    """
    Returns the full filename (without path) that matches the given partial filename within the specified directory.
    
    Args:
        partial_filename (str): Partial filename to search for.
        search_dir (str): Directory path to search in.
    
    Returns:
        str: Full filename if a match is found, else None.
    """
    try:
        for root, dirs, files in os.walk(search_dir):
            for file in files:
                if partial_filename in file:
                    return file  # Return just the filename
        print("get_full_filename function did NOT return a value")
        return None  # Return None if no match is found
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def remove_second_colon(para_text, paragraph):
    # This regex will match ": (כח):" and remove the second colon at the end
    para_text = re.sub(r"(:\((\D+)\)‪):", r" (\2)‪:", para_text)
    
    # Now we return the modified paragraph with the second colon removed
    new_paragraph = paragraph
    new_paragraph.clear()  # Clear previous runs
    new_paragraph.add_run(para_text)  # Add the modified text without the second colon
    
    return new_paragraph

def weave_torah_files(parasha_name, hebrew_file_path, english_file_path, output_file_path):
    """
    Combine Hebrew and English texts from Word documents into a single Word file.
    Preserve original formatting of Hebrew and English documents.

    Args:
        hebrew_file_path (str): Path to the Hebrew Word document.
        english_file_path (str): Path to the English Word document.
        output_file_path (str): Path for the combined Word document output.
    """
    # Load Hebrew and English Word files
    hebrew_doc = Document(hebrew_file_path)
    english_doc = Document(english_file_path)

    # Prepare the output Word document
    output_doc = Document()

    # Set narrow margins
    section = output_doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    hebrew_index = 0
    english_index = 0

    # Process paragraphs and align Hebrew and English text
    while hebrew_index < len(hebrew_doc.paragraphs):
        # Add the Hebrew paragraph with formatting
        hebrew_paragraph = hebrew_doc.paragraphs[hebrew_index]
        output_paragraph = output_doc.add_paragraph()

        # Set the Hebrew font and size
        for run in hebrew_paragraph.runs:
            run_new = output_paragraph.add_run(run.text)
            run_new.bold = run.bold
            run_new.italic = run.italic
            run_new.underline = run.underline
            run_new.font.name = utils.DOCX_HEBREW_FONT
            run_new.font.size = Pt(utils.FONT_SIZE_HEB)

        # Align Hebrew text to the right (Right-to-left alignment)
        output_paragraph.alignment = 1
        output_paragraph.paragraph_format.bidi = True

        hebrew_index += 1

        # Add the corresponding English paragraph with formatting
        if english_index < len(english_doc.paragraphs):
            english_paragraph = english_doc.paragraphs[english_index]
            output_paragraph = output_doc.add_paragraph()

            # Set the English font and size
            for run in english_paragraph.runs:
                run_new = output_paragraph.add_run(run.text)
                run_new.bold = run.bold
                run_new.italic = run.italic
                run_new.underline = run.underline
                run_new.font.name = utils.DOCX_ENGLISH_FONT
                run_new.font.size = Pt(utils.FONT_SIZE_ENG)

            # Align English text to the left (Left-to-right alignment)
            output_paragraph.alignment = 0
            output_paragraph.paragraph_format.bidi = False

            english_index += 1

    # Check if Hebrew paragraphs exist and prepare the first line for the file name
    if hebrew_doc.paragraphs:
        first_line_hebrew = hebrew_doc.paragraphs[0].text.strip()
        safe_first_line_hebrew = "".join([c if c.isalnum() or c.isspace() else "_" for c in first_line_hebrew])
    else:
        safe_first_line_hebrew = "combined"

    # Ensure the output path exists
    os.makedirs(output_file_path, exist_ok=True)

    # Define the output file path using the cleaned Hebrew text
    docx_name = f"{parasha_name}_{safe_first_line_hebrew}.docx"
    docx_name = clean_hebrew_filename(docx_name)
    final_path_result = os.path.join(output_file_path, docx_name)

    # Remove the existing file if it already exists
    if os.path.exists(final_path_result):
        os.remove(final_path_result)
        print(f"Existing file found and removed: {final_path_result}")

    # Save the combined document
    output_doc.save(final_path_result)
    print(f"Combined document saved as: {final_path_result}")
    
    return final_path_result

def add_notes_to_verses(file_path):
    """
    Adds a notes section to each verse in a .docx file while preserving formatting 
    and replaces the original file. The output file has narrow margins.

    Args:
        file_path (str): Path to the .docx file to process.
    """
    # Load the document
    doc = Document(file_path)
    
    # Extract the header (assumes it's the first paragraph with the chapter name)
    header_text = None
    for para in doc.paragraphs:
        if "Chapter" in para.text:
            header_text = para.text.strip()
            break
    
    if not header_text:
        raise ValueError("No header containing 'Chapter' found in the document.")
    
    # Create a new document to rewrite the content
    new_doc = Document()

    # Set narrow margins for the new document
    section = new_doc.sections[0]
    section.left_margin = Inches(0.5)   # 0.5 inches left margin
    section.right_margin = Inches(0.5)  # 0.5 inches right margin
    section.top_margin = Inches(0.5)    # 0.5 inches top margin
    section.bottom_margin = Inches(0.5) # 0.5 inches bottom margin

    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        # Copy the paragraph while preserving formatting
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.alignment = para.alignment  # Preserve alignment
        for run in para.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
        
        # Check if the paragraph is a verse
        if para_text.startswith("Verse"):
            # Extract the verse number using regex
            match = re.match(r"Verse (\d+)", para_text)
            if match:
                verse_number = match.group(1)
                # Add the notes paragraph for the verse
                notes_paragraph = new_doc.add_paragraph()
                notes = f"[notes]( {header_text} Verse {verse_number} )[end_notes]"
                notes_paragraph.alignment = 0  # Align notes to the left (customize if needed)
                notes_run = notes_paragraph.add_run(notes)
                # Optional: Style the notes text (customize if needed)
                notes_run.italic = True
                notes_run.font.color.rgb = RGBColor(211, 211, 211)  # Set the notes text to light grey
    
    # Delete the original file if it exists
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Existing file deleted: {file_path}")
    
    # Save the reformatted content back to the original file
    new_doc.save(file_path)
    print(f"Formatted document saved as: {file_path}")

def clean_hebrew_filename(filename):
    """
    Clean a Hebrew filename string to a consistent format.
    
    Args:
        filename (str): The original filename string.
        
    Returns:
        str: The cleaned filename string.
    """
    # Replace spaces with underscores
    cleaned_string = re.sub(r"\s+", "_", filename)
    # Replace multiple underscores with a single underscore
    cleaned_string = re.sub(r"_{2,}", "_", cleaned_string)
    # Ensure proper format for "Chapter"
    cleaned_string = re.sub(r"_Chapter_", " Chapter_", cleaned_string)
    # Ensure proper format for "Verses"
    cleaned_string = re.sub(r"_Verses_", " Verses_", cleaned_string)
    # Remove extra underscores after "Chapter" and "Verses"
    cleaned_string = re.sub(r"(?<=Chapter)_+", "_", cleaned_string)
    cleaned_string = re.sub(r"(?<=Verses)_+", "_", cleaned_string)
    # Remove trailing underscores
    cleaned_string = re.sub(r"_+$", "", cleaned_string)
    # Handle dangling underscores before ".docx"
    cleaned_string = re.sub(r"_\.docx$", ".docx", cleaned_string)
    
    return cleaned_string

def format_docx_file(file_path):
    """
    Formats the content of a .docx file, applies consistent formatting while preserving
    the content as-is, and replaces the original file.

    Args:
        file_path (str): Path to the .docx file to process.
    """
    # Load the document
    doc = Document(file_path)
    
    # Create a new document to rewrite the content
    new_doc = Document()

    # Set narrow margins for the new document
    section = new_doc.sections[0]
    section.left_margin = Inches(0.5)   # 0.5 inches left margin
    section.right_margin = Inches(0.5)  # 0.5 inches right margin
    section.top_margin = Inches(0.5)    # 0.5 inches top margin
    section.bottom_margin = Inches(0.5) # 0.5 inches bottom margin

    # Regex to detect Hebrew characters
    hebrew_range = r"[\u0590-\u05FF]"

    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        # Create a new paragraph in the formatted document
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.alignment = para.alignment  # Preserve alignment
        
        # Loop through the runs in the original paragraph and reapply formatting
        for run in para.runs:
            new_run = new_paragraph.add_run(run.text)
            
            # Preserve original run formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb

            # Apply consistent font and size based on content
            if re.search(hebrew_range, run.text):  # If the text contains Hebrew characters
                new_run.font.name = utils.DOCX_HEBREW_FONT
                new_run.font.size = Pt(utils.FONT_SIZE_HEB)
            else:  # Non-Hebrew text
                new_run.font.name = utils.DOCX_ENGLISH_FONT
                new_run.font.size = Pt(utils.FONT_SIZE_ENG)

    # Delete the original file if it exists
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"Existing file deleted: {file_path}")
    
    # Save the reformatted content back to the original file
    new_doc.save(file_path)
    print(f"Formatted document saved as: {file_path}")

def get_user_input():
    """Ask the user whether to add notes to verses."""
    return input("\nWould you like to add notes to the verses? (yes/no): ").strip().lower()


def get_parasha_details(is_hebrew=True):
    """Retrieve and print Parasha details for either Hebrew or English."""
    parasha_path = utils.load_data(utils.PARASHOT_NOW_HEB if is_hebrew else utils.PARASHOT_NOW, return_path_only=True)
    details = utils.get_parasha_details_heb2(parasha_path) if is_hebrew else utils.get_parasha_details_heb(parasha_path)

    if details:
        first_parasha = details[0]
        print("\nParasha Details (Hebrew)" if is_hebrew else "\nParasha Details (English)")
        for key, value in first_parasha.items():
            print(f"{key.replace('_', ' ').title()}:\t {value}")
        return first_parasha
    else:
        print("No parasha details found.")
        return None


def get_file_paths(parasha_name, book_name, start_chapter, num_parasha=None, is_hebrew=True):
    """Construct and return file paths for Hebrew and English texts."""
    folder_path = utils.load_tanakh_path(utils.HEB_DOCX_FOLDER if is_hebrew else utils.ENG_DOCX_FOLDER)
    end_folder_name = f"{num_parasha}_{parasha_name}" if is_hebrew else parasha_name
    folder_path = os.path.join(folder_path, end_folder_name)

    filename = f"{book_name}_CH_{start_chapter}" if is_hebrew else f"{book_name}_{start_chapter}.docx"
    filename = get_full_filename(filename, folder_path) if is_hebrew else filename + ".docx"

    print(f"\n{'Hebrew' if is_hebrew else 'English'} filename:\t {filename}")
    print(f"{'Hebrew' if is_hebrew else 'English'} folder path:\t {folder_path}")

    return os.path.join(folder_path, filename)

def weave_and_save_files(parasha_name_heb, heb_file, english_file, output_folder_path):
    """Weave Hebrew and English files together and return the output path."""
    final_output = weave_torah_files(parasha_name_heb, heb_file, english_file, output_folder_path)
    print(f"\nFinal woven document:\t {final_output}")
    return final_output


def process_document(final_output, add_notes):
    """Process the document by adding notes (if needed), removing colons, and formatting."""
    if add_notes == 'yes':
        add_notes_to_verses(final_output)
        print("\nNotes have been added to the verses.")

    doc = Document(final_output)
    
    # Remove second Hebrew verse number colon
    for para in doc.paragraphs:
        para_text = para.text.strip()
        remove_second_colon(para_text, para)

    doc.save(final_output)
    format_docx_file(final_output)
    print(f"\nDocument saved and formatted: {final_output}")

def contains_hebrew(text):
    """
    Check if the provided text contains Hebrew characters.

    Parameters:
    - text (str): The text to check.

    Returns:
    - bool: True if the text contains Hebrew characters, False otherwise.
    """
    # Hebrew characters range from Unicode 0x0590 to 0x05FF
    return any("\u0590" <= char <= "\u05FF" for char in text)

def apply_rtl_alignment(paragraph):
    """
    Apply right-to-left alignment to the paragraph and ensure it is set at the XML level.
    
    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph to format.
    """
    # Align paragraph text to the right
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Ensure the XML-level setting for RTL is applied
    paragraph._p.set(qn('w:bidi'), '1')

def format_hebrew_paragraph(document):
    """
    Format Hebrew text in a Word document:
    - Align the text to the right (RTL).
    - Move the verse number to the beginning of the paragraph.
    - Apply specific fonts and font sizes.

    Parameters:
    - document (docx.Document): The document to be formatted.
    """
    
    for paragraph in document.paragraphs:
        # Step 1: Check if the paragraph contains any Hebrew characters
        if contains_hebrew(paragraph.text):
            # Step 2: Apply right-to-left (RTL) alignment
            apply_rtl_alignment(paragraph)
            
            # Step 3: Move the verse number (if present) to the beginning of the paragraph
            paragraph.text = move_verse_number_to_start(paragraph.text)
            
            # Step 4: Apply the Hebrew font and size to the paragraph
            apply_hebrew_font(paragraph)

def update_second_line(document, parasha_name):
    """
    Update the second line of the document by removing everything before 'Chapter' 
    and appending the parasha_name.
    Parameters:
    - document (docx.Document): The document object to process.
    - parasha_name (str): The name of the parasha to append to the second line.
    """
    # Ensure the document has at least two paragraphs
    if len(document.paragraphs) >= 2:
        second_paragraph = document.paragraphs[1]
        if "Chapter" in second_paragraph.text:
            # Find the position of the word "Chapter" and update the text
            chapter_index = second_paragraph.text.find("Chapter")
            updated_text = second_paragraph.text[chapter_index:] + f" - {parasha_name}"
            second_paragraph.text = updated_text

def format_eng_paragraph(paragraph):
    """
    Process the paragraph to replace "::" with ":" while preserving formatting.

    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph object to process.
    """
    # Combine all runs' text into a single string
    para_text = "".join(run.text for run in paragraph.runs)

    # Replace "::" with ":" in the combined text
    updated_text = remove_second_colon_eng(para_text)

    # If changes are needed, clear existing runs and re-add the updated text
    if para_text != updated_text:
        for run in paragraph.runs:
            run.text = ""  # Clear existing text
        paragraph.add_run(updated_text)  # Add updated text back

def move_verse_number_to_start(text):
    """
    Move the verse number (if present) to the start of the paragraph.
    
    Parameters:
    - text (str): The paragraph text.
    
    Returns:
    - str: The modified text with the verse number moved to the beginning.
    """
    #print(f"Original text: {text}")  # Debug: Show the original text
    
    # Regular expression to match verse numbers inside parentheses, including special characters like ‪
    match = re.search(r"‪\s?\(([^)]+)\)\s?‪", text)  # Match "(כח)", "(כט)", "(לא)", etc. allowing for spaces or special characters
    
    if match:
        # Extract the verse number (e.g., "כח", "כט")
        verse_number = match.group(1)
        #print(f"Verse number found: {verse_number}")  # Debug: Show the found verse number
        
        # Remove the verse number from its original position
        modified_text = re.sub(r"‪\s?\(.*?\)\s?‪", "", text).strip()  # Strip any unwanted spaces
        
        # Remove any trailing colon (:) from the text
        if modified_text.endswith(":"):
            modified_text = modified_text[:-1]
        
        # Build the new text with the verse number at the start
        modified_text = modified_text + "\u200F"
        verse_number = f"\u200F{verse_number}  \u200F"
        new_text = verse_number + " \u200F" + modified_text + ":\u200F"
        #print(f"New text after modification: {new_text}")  # Debug: Show the modified text
        
        return new_text
    else:
        print("No verse number found.")  # Debug: Inform when no verse number is found
        # If no verse number is found, return the original text
        return text

def apply_hebrew_font(paragraph):
    """
    Apply the specific Hebrew font and size to the paragraph text.
    
    Parameters:
    - paragraph (docx.text.Paragraph): The paragraph to format.
    """
    for run in paragraph.runs:
        # Apply Hebrew font and size
        run.font.name = utils.DOCX_HEBREW_FONT
        run.font.size = Pt(utils.FONT_SIZE_HEB)

def remove_second_colon_eng(para_text):
    """
    Replace all occurrences of "::" with ":" in the given text.

    Parameters:
    - para_text (str): The text of the paragraph.

    Returns:
    - str: The updated text with "::" replaced by ":".
    """
    # Replace all occurrences of "::" with ":"
    return re.sub(r"::", r":", para_text)

if __name__ == "__main__":
    add_notes = get_user_input()

    # Retrieve Parasha details
    parasha_details = get_parasha_details(is_hebrew=False)
    parasha_details_heb = get_parasha_details(is_hebrew=True)

    if parasha_details and parasha_details_heb:
        # Extract necessary details
        parasha_name, book_name, start_chapter = (
            parasha_details["parasha_name"], 
            parasha_details["book_name"], 
            parasha_details["start_chapter"]
        )
        
        parasha_name_heb, book_name_heb, start_chapter_heb, end_chapter_heb, num_parasha = (
            parasha_details_heb["parasha_name"],
            parasha_details_heb["book_name"],
            parasha_details_heb["start_chapter"],
            parasha_details_heb["end_chapter"],
            parasha_details_heb["num_parasha"],
        )

        # Iterate through the chapter range
        for chapter in range(start_chapter_heb, end_chapter_heb + 1):
            # Get file paths for each chapter
            english_file = get_file_paths(parasha_name, book_name, chapter, is_hebrew=False)
            heb_file = get_file_paths(parasha_name_heb, book_name_heb, chapter, num_parasha, is_hebrew=True)

            # Weave and save document
            final_output = weave_and_save_files(parasha_name_heb, heb_file, english_file, utils.load_tanakh_path(utils.OUTPUT_DOCX_FOLDER))

            # Process final document
            process_document(final_output, add_notes)

            doc = Document(final_output)

            # Setp 2: Update the line after the header
            update_second_line(doc,parasha_name_heb)

            # Step 3: Format the Hebrew text in Word
            #    - Align the Heb text to the right (RTL).
            #    - Move the verse number to the beginning of the paragraph.
            #    - Apply specific fonts and font sizes.
            format_hebrew_paragraph(doc)

            # Step 4: Format the Eng text in Word
            #    - replace "::" with ":" in the Engs
            for para in doc.paragraphs:
                format_eng_paragraph(para)

            # Step 5: Save the updated document
            doc.save(final_output)
            print("File has been updated and saved.")
