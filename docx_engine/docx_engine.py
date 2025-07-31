from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import sys
import os

# -------------------------
# Bootstrapping Dependencies
# -------------------------
# Get the absolute path to the *parent* of the current file's directory
BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent

# Folders in the root directory that contain modules
DEPENDENCY_DIRS = [
    BASE_DIR / "excel_engine",
    PROJECT_ROOT / "utils"
]

# Add each dependency directory to sys.path if not already added
for path in DEPENDENCY_DIRS:
    path_str = str(path)
    if path_str not in sys.path:
        sys.path.append(path_str)

import utils                      # utils directory
import TanachXML_engine           # xml_engine directory
import excel_engine               # excel_engine directory

# DOCX Constants
DOCX_HEBREW_FONT = "Frank Ruehl"  # Use Frank Ruehl for Hebrew text on Word
DOCX_ENGLISH_FONT = "Times New Roman"  # Use Times New Roman for English text
FONT_SIZE_HEB = 16  # Font size in points
FONT_SIZE_ENG = 12  # Font size in points
MARGIN_SIZE = Pt(12)  # Margin size in points

def create_docx_with_header(eng_str: str, heb_str: str, file_path: str, file_name: str):
    """
    Create a DOCX document with a styled bilingual header.

    Args:
        eng_str (str): English header text.
        heb_str (str): Hebrew header text.
        file_path (str): Folder path to save the document.
        file_name (str): Name of the DOCX file to save.

    Returns:
        Document: A python-docx Document object with the header created.
    """
    doc = Document()

    # Add Hebrew header
    heb_para = doc.add_paragraph()
    heb_run = heb_para.add_run(heb_str)
    heb_run.font.name = DOCX_HEBREW_FONT
    heb_run.font.size = Pt(FONT_SIZE_HEB)
    heb_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add English header
    eng_para = doc.add_paragraph()
    eng_run = eng_para.add_run(eng_str)
    eng_run.font.name = DOCX_ENGLISH_FONT
    eng_run.font.size = Pt(FONT_SIZE_ENG)
    eng_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = MARGIN_SIZE
        section.bottom_margin = MARGIN_SIZE
        section.left_margin = MARGIN_SIZE
        section.right_margin = MARGIN_SIZE

    # Save file
    os.makedirs(file_path, exist_ok=True)
    full_path = os.path.join(file_path, file_name)
    doc.save(full_path)

    return doc

def append_paragraph_to_docx(doc: Document, text: str, is_hebrew: bool = False):
    """
    Appends a paragraph with specified text to a DOCX Document object.

    Args:
        doc (Document): An existing python-docx Document object.
        text (str): Text content to append.
        is_hebrew (bool): Whether the text is Hebrew (sets font/alignment accordingly).
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    
    if is_hebrew:
        run.font.name = DOCX_HEBREW_FONT
        run.font.size = Pt(FONT_SIZE_HEB)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        run.font.name = DOCX_ENGLISH_FONT
        run.font.size = Pt(FONT_SIZE_ENG)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def get_metsudah_ch_docx(hc_book: str, hc_book_heb: str, hc_chapter: int):
    """
    Export a full chapter of Torah to a DOCX file with Hebrew and English verses.

    Args:
        hc_book (str): Book name in English (e.g., "Genesis")
        hc_chapter (int): Chapter number (1-based)
    """

    # XML Book Name
    hc_book_xml = str(hc_book) + ".xml"

    # Get number of verses in the chapter
    num_verses = utils.get_torah_ch_verse_num(hc_book, hc_chapter)

    # Prepare DOCX header and file path
    header = f"{hc_book} Chapter {hc_chapter}"
    heb_text = f"תּוֹרָה - סֵפֶר {hc_book_heb}"
    file_name = f"{hc_book}_Ch_{hc_chapter}.docx"

    doc = create_docx_with_header(header, heb_text, utils.METSUDAH_DOCX_ENG_OUTPUT, file_name)

    # Loop through each verse
    for hc_verse in range(1, num_verses + 1):
        # Get Hebrew verse with verse number
        heb_verse_num = utils.get_hebrew_verse_num(hc_verse, utils.H_VERSE_NUM_JSON)
        heb_verse = TanachXML_engine.get_verse(utils.HEB_TORAH_BOOK_DATA_XML, hc_book_xml, hc_chapter, hc_verse)
        heb_verse = " ".join(heb_verse)
        heb_string = f"{heb_verse_num}   {heb_verse}"
        append_paragraph_to_docx(doc, heb_string, is_hebrew=True)

        # Get English verse from Excel
        row = hc_verse + 1  # +1 to skip header
        sheet = f"{hc_book} CH{hc_chapter}"
        xlsx_metsudah_eng = f"{hc_book}.xlsx"
        eng_metsudah_xlsx_path = utils.METSUDAH_XLSX_ENG_FILES / xlsx_metsudah_eng
        a, b = excel_engine.get_excel_row_ab(eng_metsudah_xlsx_path, sheet, row)
        eng_string = f"{a} {b}"
        append_paragraph_to_docx(doc, eng_string, is_hebrew=False)

    # Save file
    full_path = os.path.join(utils.METSUDAH_DOCX_ENG_OUTPUT, file_name)
    doc.save(full_path)
    print(f"[INFO] Saved chapter to: {full_path}")

# Example usage:
# create_docx_with_header("Genesis - Chapter 1", "\u05d1\u05e8\u05d0\u05e9\u05d9\u05ea \u05e4\u05e8\u05e7 \u05d0", "/path/to/folder", "output.docx")
